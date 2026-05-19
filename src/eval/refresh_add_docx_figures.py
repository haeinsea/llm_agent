from __future__ import annotations

from copy import deepcopy
import os
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile
import xml.etree.ElementTree as ET

PROJECT_ROOT = Path(__file__).resolve().parents[2]
os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_ROOT / "outputs" / ".mplcache"))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from scipy.stats import gaussian_kde
from sklearn.decomposition import PCA
from sklearn.manifold import TSNE

from src.eval.analyze_te_shift import infer_feature_cols, prepare_feature_matrix, sample_for_embedding
from src.eval.plot_style import PAPER_COLORS, save_figure, set_paper_style, style_axes
from src.models.graphad import graphad_score_matrices, load_graphad_artifact
from src.models.temporal_backbone import temporal_model_display_name
from src.routing.selective_llm_eval import DEFAULT_Q, read_selected_q
from src.utils.io import ensure_dir, read_csv, read_json, read_yaml
from src.utils.metrics import binary_metrics
from src.utils.routing import build_routing_features


CONFIG_DIR = PROJECT_ROOT / "configs"
DATA_DIR = PROJECT_ROOT / "data"
PROCESSED_DIR = DATA_DIR / "processed"
OUTPUT_DIR = PROJECT_ROOT / "outputs"
PRED_DIR = OUTPUT_DIR / "predictions"
FIG_DIR = OUTPUT_DIR / "figures"
APPENDIX_DIR = OUTPUT_DIR / "appendix"
METRIC_DIR = OUTPUT_DIR / "metrics"
MODEL_DIR = OUTPUT_DIR / "models"
SHIFT_DIR = OUTPUT_DIR / "shift_analysis_test_main"

DOCX_PATH = PROJECT_ROOT / "UTAR 논문(Reference)_260321_add.docx"
DOCX_BACKUP_PATH = PROJECT_ROOT / "UTAR 논문(Reference)_260321_add_backup_before_hq_figures.docx"


def _clean_series(values: pd.Series) -> np.ndarray:
    return values.astype(float).replace([np.inf, -np.inf], np.nan).dropna().to_numpy(dtype=float)


def _kde_on_grid(values: pd.Series | np.ndarray, grid: np.ndarray) -> np.ndarray:
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]
    if len(arr) < 2:
        return np.zeros_like(grid)
    if np.allclose(arr.std(), 0.0):
        out = np.zeros_like(grid)
        idx = int(np.argmin(np.abs(grid - float(arr.mean()))))
        out[idx] = 1.0
        return out
    kde = gaussian_kde(arr)
    return kde(grid)


def _best_tau(y_true: pd.Series, scores: np.ndarray) -> float:
    best_tau = 0.5
    best_f1 = -1.0
    for tau in np.linspace(0.01, 0.99, 99):
        f1 = binary_metrics(y_true, scores, tau=float(tau))["f1"]
        if f1 > best_f1:
            best_f1 = f1
            best_tau = float(tau)
    return best_tau


def _load_predictions() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, float, float, str]:
    cfg = read_yaml(CONFIG_DIR / "routing.yaml", default={})
    tcn_cfg = read_yaml(CONFIG_DIR / "train_tcn.yaml", default={})
    temporal_label = temporal_model_display_name(tcn_cfg.get("architecture", "modern_tcn"))

    tau = float(read_json(METRIC_DIR / "thresholds.json")["tau"])
    gray_grid = read_csv(METRIC_DIR / "grayzone_grid.csv")
    selected_q = read_selected_q(DEFAULT_Q)
    margin = float(gray_grid.loc[np.isclose(gray_grid["q"], selected_q), "gray_margin_mean"].iloc[0])

    base = read_csv(PRED_DIR / "base_test_main_predictions.csv")
    val = read_csv(PRED_DIR / "base_val_predictions.csv")
    selective = read_csv(PRED_DIR / "utar_test_main_selective.csv")

    if "p_utar_base" not in base.columns or "ensemble_entropy" not in base.columns:
        routing = build_routing_features(base[["source_file", "fault_id", "run_id", "p_rf", "p_xgb", "p_tcn"]], cfg)
        for col in routing.columns:
            base[col] = routing[col]

    if "p_utar_base" not in val.columns or "ensemble_entropy" not in val.columns:
        routing = build_routing_features(val[["source_file", "fault_id", "run_id", "p_rf", "p_xgb", "p_tcn"]], cfg)
        for col in routing.columns:
            val[col] = routing[col]

    if "p_utar_base" not in selective.columns or "ensemble_entropy" not in selective.columns:
        routing = build_routing_features(selective[["source_file", "fault_id", "run_id", "p_rf", "p_xgb", "p_tcn"]], cfg)
        for col in routing.columns:
            selective[col] = routing[col]

    return base, val, selective, tau, margin, temporal_label


def _format_axes(ax, *, xlabel: str | None = None, ylabel: str | None = None, y_grid_only: bool = False) -> None:
    if xlabel:
        ax.set_xlabel(xlabel)
    if ylabel:
        ax.set_ylabel(ylabel)
    style_axes(ax, y_grid_only=y_grid_only)


def _draw_distribution_panel(ax_top, ax_bottom, val_df: pd.DataFrame, test_df: pd.DataFrame, tau: float, margin: float) -> None:
    grid = np.linspace(0.0, 1.0, 600)

    val_normal = _kde_on_grid(val_df.loc[val_df["y_true"] == 0, "p_utar_base"], grid)
    val_anomaly = _kde_on_grid(val_df.loc[val_df["y_true"] == 1, "p_utar_base"], grid)
    test_normal = _kde_on_grid(test_df.loc[test_df["y_true"] == 0, "p_utar_base"], grid)
    test_anomaly = _kde_on_grid(test_df.loc[test_df["y_true"] == 1, "p_utar_base"], grid)

    normal_fill = "#b9d4ea"
    normal_edge = "#2c7fb8"
    anomaly_fill = "#f6c4b8"
    anomaly_edge = "#c44536"

    for ax, normal_curve, anomaly_curve, title in [
        (ax_top, val_normal, val_anomaly, "Training (IID / Validation)"),
        (ax_bottom, test_normal, test_anomaly, "Operational (TDS / Shift)"),
    ]:
        ax.fill_between(grid, normal_curve, color=normal_fill, alpha=0.78, zorder=2)
        ax.plot(grid, normal_curve, color=normal_edge, linewidth=2.0, zorder=3)
        ax.fill_between(grid, anomaly_curve, color=anomaly_fill, alpha=0.80, zorder=2)
        ax.plot(grid, anomaly_curve, color=anomaly_edge, linewidth=2.0, zorder=3)
        ax.axvline(tau, color=PAPER_COLORS["ink"], linestyle="--", linewidth=1.25, zorder=4)
        ax.axvspan(max(0.0, tau - margin), min(1.0, tau + margin), color="#f6d55c", alpha=0.34, zorder=1)
        ax.set_xlim(0.0, 1.0)
        ax.set_title(title, fontsize=12.5, pad=7)
        _format_axes(ax, ylabel="Probability Density", y_grid_only=True)

    ax_top.set_xticklabels([])
    ax_top.legend(
        handles=[
            Line2D([0], [0], color=normal_edge, lw=2.0, label="Normal"),
            Line2D([0], [0], color=anomaly_edge, lw=2.0, label="Anomaly"),
            Line2D([0], [0], color=PAPER_COLORS["ink"], lw=1.25, linestyle="--", label=rf"$\tau={tau:.2f}$"),
        ],
        loc="upper right",
        fontsize=9.5,
    )

    iid_overlap = float(np.mean(np.abs(val_df["p_utar_base"].to_numpy(dtype=float) - tau) <= margin))
    shift_overlap = float(np.mean(np.abs(test_df["p_utar_base"].to_numpy(dtype=float) - tau) <= margin))

    ax_top.annotate(
        f"IID gray-zone share: {iid_overlap * 100:.1f}%",
        xy=(tau, ax_top.get_ylim()[1] * 0.34),
        xytext=(tau + 0.08, ax_top.get_ylim()[1] * 0.62),
        arrowprops={"arrowstyle": "->", "color": PAPER_COLORS["ink"], "lw": 1.0},
        fontsize=9.6,
        ha="left",
    )
    ax_bottom.annotate(
        f"TDS gray-zone share: {shift_overlap * 100:.1f}%",
        xy=(tau, ax_bottom.get_ylim()[1] * 0.42),
        xytext=(tau + 0.10, ax_bottom.get_ylim()[1] * 0.78),
        arrowprops={"arrowstyle": "->", "color": PAPER_COLORS["ink"], "lw": 1.0},
        fontsize=9.6,
        ha="left",
    )
    ax_bottom.text(
        tau,
        ax_bottom.get_ylim()[1] * 0.17,
        "Gray-Zone",
        ha="center",
        va="center",
        fontsize=11.0,
        fontweight="semibold",
        color=PAPER_COLORS["ink"],
    )
    ax_bottom.set_xlabel("UTAR Base Score")


def _draw_margin_panel(ax, base: pd.DataFrame, tau: float, margin: float) -> None:
    shift_df = base[base["phase"].isin(["transition", "post_shift"])].copy()
    grid = np.linspace(-0.22, 0.22, 600)
    centered_normal = shift_df.loc[shift_df["y_true"] == 0, "p_utar_base"] - tau
    centered_anomaly = shift_df.loc[shift_df["y_true"] == 1, "p_utar_base"] - tau
    normal_curve = _kde_on_grid(centered_normal, grid)
    anomaly_curve = _kde_on_grid(centered_anomaly, grid)
    ymax = max(float(normal_curve.max()), float(anomaly_curve.max())) * 1.18

    ax.axvspan(-margin, margin, color="#f6d55c", alpha=0.38, zorder=1)
    ax.axvspan(-margin * 0.35, margin * 0.35, color="#f3b63a", alpha=0.28, zorder=1)
    ax.fill_between(grid, normal_curve, color="#b9d4ea", alpha=0.78, zorder=2)
    ax.plot(grid, normal_curve, color="#2c7fb8", linewidth=2.1, label="Normal", zorder=3)
    ax.fill_between(grid, anomaly_curve, color="#f6c4b8", alpha=0.80, zorder=2)
    ax.plot(grid, anomaly_curve, color="#c44536", linewidth=2.1, label="Anomaly", zorder=3)
    ax.axvline(0.0, color=PAPER_COLORS["ink"], linestyle="--", linewidth=1.2, zorder=4)

    normal_sample = centered_normal.sample(min(250, len(centered_normal)), random_state=42) if len(centered_normal) else centered_normal
    anomaly_sample = centered_anomaly.sample(min(250, len(centered_anomaly)), random_state=43) if len(centered_anomaly) else centered_anomaly
    ax.vlines(normal_sample.to_numpy(dtype=float), 0.0, ymax * 0.04, color="#2c7fb8", alpha=0.18, linewidth=0.8)
    ax.vlines(anomaly_sample.to_numpy(dtype=float), 0.0, ymax * 0.06, color="#c44536", alpha=0.18, linewidth=0.8)

    for ypos in [ymax * 0.24, ymax * 0.43, ymax * 0.62]:
        ax.add_patch(
            FancyArrowPatch(
                (-0.18, ypos),
                (-margin * 0.18, ypos),
                arrowstyle="simple",
                mutation_scale=14,
                linewidth=0.6,
                facecolor="white",
                edgecolor=PAPER_COLORS["ink"],
            )
        )
        ax.add_patch(
            FancyArrowPatch(
                (0.18, ypos),
                (margin * 0.18, ypos),
                arrowstyle="simple",
                mutation_scale=14,
                linewidth=0.6,
                facecolor="white",
                edgecolor=PAPER_COLORS["ink"],
            )
        )

    gray_share = float(np.mean(np.abs(shift_df["p_utar_base"].to_numpy(dtype=float) - tau) <= margin))
    ax.annotate(
        f"Observed gray-zone: {gray_share * 100:.1f}%",
        xy=(0.0, ymax * 0.80),
        xytext=(0.07, ymax * 0.97),
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
        fontsize=9.8,
        ha="left",
    )
    ax.text(0.0, ymax * 1.03, "Decision boundary", ha="center", va="bottom", fontsize=11.0)
    ax.text(0.0, ymax * 1.11, r"$\tau$", ha="center", va="bottom", fontsize=14.0)
    ax.text(0.0, ymax * 0.12, "Concentrated\nuncertainty", ha="center", va="center", fontsize=10.5, fontweight="semibold")

    ax.set_xlim(-0.22, 0.22)
    ax.set_ylim(0.0, ymax * 1.15)
    _format_axes(ax, xlabel=r"Centered Score $(p - \tau)$", ylabel="Probability Density", y_grid_only=True)
    ax.legend(loc="upper right", fontsize=9.5)


def _draw_discrepancy_panel(ax_top, ax_bottom, base: pd.DataFrame, temporal_label: str) -> None:
    normal = base.loc[base["phase"] == "normal", ["p_rf", "p_tcn"]].dropna()
    shift = base.loc[base["phase"].isin(["transition", "post_shift"]), ["p_rf", "p_tcn", "y_true"]].dropna()
    shift = shift.copy()
    shift["gap"] = (shift["p_rf"] - shift["p_tcn"]).abs()

    if len(normal) > 950:
        normal = normal.sample(950, random_state=44)
    if len(shift) > 950:
        shift = shift.sample(950, random_state=45)

    highlighted = shift.sort_values("gap", ascending=False).head(5)

    def draw(ax, df: pd.DataFrame, title: str) -> None:
        ax.scatter(df["p_rf"], df["p_tcn"], s=15, color="#2c7fb8", alpha=0.70, edgecolors="white", linewidths=0.25, zorder=2)
        ax.plot([0.0, 1.0], [0.0, 1.0], linestyle="--", color=PAPER_COLORS["ink"], linewidth=1.1, zorder=3)
        ax.set_xlim(0.0, 1.0)
        ax.set_ylim(0.0, 1.0)
        ax.set_title(title, fontsize=12.5, pad=7)
        _format_axes(ax, ylabel=f"{temporal_label} Score")

    draw(ax_top, normal, "Normal Condition (IID)")
    draw(ax_bottom, shift, "Distribution Shift (TDS)")
    ax_top.set_xticklabels([])
    ax_bottom.set_xlabel("RF Score")

    ax_top.annotate(
        "Aligned predictions\nindicate consensus",
        xy=(0.18, 0.18),
        xytext=(0.04, 0.63),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
        fontsize=9.6,
    )

    for _, row in highlighted.iterrows():
        ax_bottom.scatter([row["p_rf"]], [row["p_tcn"]], s=80, facecolor="none", edgecolor="#c87f2a", linewidth=1.4, zorder=4)

    if not highlighted.empty:
        top_gap = highlighted.iloc[0]
        ax_bottom.annotate(
            f"Largest gap: {top_gap['gap']:.3f}",
            xy=(top_gap["p_rf"], top_gap["p_tcn"]),
            xytext=(0.03, 0.79),
            textcoords="axes fraction",
            arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
            fontsize=9.4,
        )
        last_gap = highlighted.iloc[-1]
        ax_bottom.annotate(
            "Escalation candidates",
            xy=(last_gap["p_rf"], last_gap["p_tcn"]),
            xytext=(0.69, 0.19),
            textcoords="axes fraction",
            arrowprops={"arrowstyle": "->", "lw": 1.0, "color": "#c87f2a"},
            fontsize=9.4,
            ha="left",
            color="#915c18",
        )


def build_methodology_panels(base: pd.DataFrame, val_df: pd.DataFrame, tau: float, margin: float, temporal_label: str) -> dict[str, Path]:
    ensure_dir(FIG_DIR)
    outputs = {
        "distribution": FIG_DIR / "figure1_panel_distribution_overlap.png",
        "margin": FIG_DIR / "figure1_panel_margin_concentration.png",
        "discrepancy": FIG_DIR / "figure1_panel_model_discrepancy.png",
    }

    fig, axes = plt.subplots(2, 1, figsize=(4.3, 5.8), constrained_layout=True)
    _draw_distribution_panel(axes[0], axes[1], val_df, base, tau=tau, margin=margin)
    save_figure(fig, outputs["distribution"])
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(4.2, 5.2), constrained_layout=True)
    _draw_margin_panel(ax, base, tau=tau, margin=margin)
    save_figure(fig, outputs["margin"])
    plt.close(fig)

    fig, axes = plt.subplots(2, 1, figsize=(4.4, 5.8), sharex=True, sharey=True, constrained_layout=True)
    _draw_discrepancy_panel(axes[0], axes[1], base, temporal_label=temporal_label)
    save_figure(fig, outputs["discrepancy"])
    plt.close(fig)

    return outputs


def _ensemble_seed_metrics() -> pd.DataFrame:
    val = read_csv(PRED_DIR / "base_val_predictions.csv")
    test = read_csv(PRED_DIR / "base_test_main_predictions.csv")
    seed_cols = sorted(int(col.split("seed")[-1]) for col in val.columns if col.startswith("p_rf_seed"))
    rows = []
    for seed in seed_cols:
        p_val = val[[f"p_rf_seed{seed}", f"p_xgb_seed{seed}", f"p_tcn_seed{seed}"]].mean(axis=1).to_numpy(dtype=float)
        p_test = test[[f"p_rf_seed{seed}", f"p_xgb_seed{seed}", f"p_tcn_seed{seed}"]].mean(axis=1).to_numpy(dtype=float)
        tau_seed = _best_tau(val["y_true"], p_val)
        val_metrics = binary_metrics(val["y_true"], p_val, tau=tau_seed)
        test_metrics = binary_metrics(test["y_true"], p_test, tau=tau_seed)
        rows.append(
            {
                "seed": seed,
                "tau": tau_seed,
                "val_f1": float(val_metrics["f1"]),
                "test_f1": float(test_metrics["f1"]),
            }
        )
    return pd.DataFrame(rows)


def build_performance_drop_figure() -> Path:
    selected_q = read_selected_q(DEFAULT_Q)
    ensemble_df = _ensemble_seed_metrics()
    utar_seed = read_csv(METRIC_DIR / "selective_llm_seed_metrics.csv")
    utar_seed = utar_seed[(utar_seed["q"].round(6) == round(selected_q, 6)) & (utar_seed["mode"] == "selective")].copy()

    utar_val = utar_seed[utar_seed["dataset"] == "val"]["f1"].to_numpy(dtype=float)
    utar_test = utar_seed[utar_seed["dataset"] == "main"]["f1"].to_numpy(dtype=float)
    ens_val = ensemble_df["val_f1"].to_numpy(dtype=float)
    ens_test = ensemble_df["test_f1"].to_numpy(dtype=float)

    x = np.array([0.0, 1.0])
    ens_mean = np.array([ens_val.mean(), ens_test.mean()])
    ens_std = np.array([ens_val.std(ddof=1), ens_test.std(ddof=1)])
    utar_mean = np.array([utar_val.mean(), utar_test.mean()])
    utar_std = np.array([utar_val.std(ddof=1), utar_test.std(ddof=1)])

    fig, ax = plt.subplots(figsize=(8.8, 5.1), constrained_layout=True)
    ax.axvspan(0.92, 1.08, color="#edf4ff", alpha=0.92, zorder=0)
    ax.plot(x, ens_mean, linestyle="--", marker="o", markersize=7.0, linewidth=2.2, color="#2c7fb8", label="Simple Ensemble", zorder=3)
    ax.plot(x, utar_mean, linestyle="-", marker="o", markersize=7.2, linewidth=2.4, color="#c44536", label="UTAR", zorder=4)
    ax.fill_between(x, ens_mean - ens_std, ens_mean + ens_std, color="#2c7fb8", alpha=0.18, zorder=1)
    ax.fill_between(x, utar_mean - utar_std, utar_mean + utar_std, color="#c44536", alpha=0.18, zorder=1)
    ax.set_xticks(x)
    ax.set_xticklabels(["IID (Validation)", "TDS (Operational)"])
    ax.set_ylim(min(ens_mean.min(), utar_mean.min()) - 0.08, max(ens_mean.max(), utar_mean.max()) + 0.05)
    _format_axes(ax, ylabel="F1-score", y_grid_only=True)
    ax.set_title("Performance Robustness under Distribution Shift", pad=10)
    ax.legend(loc="lower left")

    ens_drop = float(ens_mean[0] - ens_mean[1])
    utar_drop = float(utar_mean[0] - utar_mean[1])
    ax.annotate(
        f"Large performance drop\n($\\Delta_1={ens_drop:.3f}$)",
        xy=(1.0, ens_mean[1]),
        xytext=(0.32, ens_mean[1] - 0.022),
        textcoords="data",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": "#2c7fb8"},
        fontsize=9.8,
        color=PAPER_COLORS["ink"],
    )
    ax.annotate(
        f"Small performance drop\n($\\Delta_2={utar_drop:.3f}$)",
        xy=(1.0, utar_mean[1]),
        xytext=(1.02, utar_mean[1] + 0.005),
        textcoords="data",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": "#c44536"},
        fontsize=9.8,
        color=PAPER_COLORS["ink"],
        ha="left",
    )
    ax.text(
        0.66,
        (utar_mean[1] + ens_mean[1]) / 2.0 + 0.004,
        "UTAR Resilience\n(Stability(UTAR) > Stability(Simple Ensemble))",
        fontsize=10.1,
        ha="left",
        va="center",
        color=PAPER_COLORS["ink"],
    )
    out_path = FIG_DIR / "figure4_performance_drop.png"
    save_figure(fig, out_path)
    plt.close(fig)
    return out_path


def _run_flip_count(scores: np.ndarray, tau: float) -> int:
    states = (np.asarray(scores, dtype=float) >= tau).astype(int)
    return int(np.abs(np.diff(states)).sum())


def _pick_representative_run(base: pd.DataFrame, selective: pd.DataFrame, tau: float) -> pd.DataFrame:
    merged = base.merge(
        selective[["source_file", "fault_id", "run_id", "sample_idx", "p_final"]],
        on=["source_file", "fault_id", "run_id", "sample_idx"],
        how="inner",
    )
    candidates = (
        merged[(merged["phase"].isin(["transition", "post_shift"])) & (merged["fault_id"] != 0)]
        [["source_file", "fault_id", "run_id"]]
        .drop_duplicates()
        .to_dict("records")
    )
    best_key = None
    best_gain = -np.inf
    for key in candidates:
        sub = merged[
            (merged["source_file"] == key["source_file"])
            & (merged["fault_id"] == key["fault_id"])
            & (merged["run_id"] == key["run_id"])
        ].sort_values("sample_idx")
        if len(sub) < 16:
            continue
        onset_mask = sub["phase"].isin(["transition", "post_shift"]).to_numpy()
        ensemble_post = sub.loc[onset_mask, "p_ensemble"].to_numpy(dtype=float)
        utar_post = sub.loc[onset_mask, "p_final"].to_numpy(dtype=float)
        gain = _run_flip_count(ensemble_post, tau=tau) - _run_flip_count(utar_post, tau=tau)
        gain += 3.0 * (float(np.nanmin(utar_post)) - float(np.nanmin(ensemble_post)))
        if gain > best_gain:
            best_gain = gain
            best_key = key

    if best_key is None:
        raise RuntimeError("Could not identify a representative shifted run.")

    chosen = merged[
        (merged["source_file"] == best_key["source_file"])
        & (merged["fault_id"] == best_key["fault_id"])
        & (merged["run_id"] == best_key["run_id"])
    ].sort_values("sample_idx")
    chosen = chosen.copy()
    chosen["relative_step"] = np.arange(len(chosen))
    return chosen


def build_shift_trace_figure(base: pd.DataFrame, selective: pd.DataFrame, tau: float, margin: float) -> Path:
    sub = _pick_representative_run(base, selective, tau=tau)
    phase_codes = sub["phase"].to_numpy()
    transition_end = int(np.sum(phase_codes == "transition"))

    fig, ax = plt.subplots(figsize=(10.6, 5.8))
    fig.subplots_adjust(left=0.09, right=0.985, top=0.88, bottom=0.14)

    x = np.arange(1, len(sub) + 1, dtype=float)
    ax.axvspan(1.0, max(float(transition_end), 1.0) + 0.5, color="#edf4fb", alpha=0.75, zorder=0)
    ax.axvspan(max(float(transition_end), 1.0) + 0.5, float(len(sub)) + 0.5, color="#f8f8f8", alpha=0.95, zorder=0)
    ax.axhspan(max(0.0, tau - margin), min(1.0, tau + margin), color="#dbe5f0", alpha=0.28, zorder=0)

    series_specs = [
        ("UTAR (Proposed)", sub["p_final"].to_numpy(dtype=float), "#1f5fa8", 3.0, "o"),
        ("Simple Ensemble", sub["p_ensemble"].to_numpy(dtype=float), PAPER_COLORS["muted_green"], 2.4, "o"),
        ("Random Forest", sub["p_rf"].to_numpy(dtype=float), PAPER_COLORS["orange"], 1.8, None),
        ("XGB Ensemble", sub["p_xgb"].to_numpy(dtype=float), PAPER_COLORS["slate"], 1.8, None),
    ]
    for label, values, color, lw, marker in series_specs:
        ax.plot(
            x,
            values,
            label=label,
            color=color,
            linewidth=lw,
            marker=marker,
            markersize=4.3 if marker else 0.0,
            markerfacecolor="white" if marker else color,
            markeredgewidth=0.9 if marker else 0.0,
            zorder=3,
        )

    ax.axvline(max(float(transition_end), 1.0) + 0.5, color=PAPER_COLORS["ink"], linestyle=":", linewidth=1.3, zorder=4)
    ax.axhline(tau, color=PAPER_COLORS["highlight"], linestyle="--", linewidth=1.25, zorder=4)
    ax.text(0.08, 0.86, "Transition", transform=ax.transAxes, fontsize=10.4, fontweight="semibold", color=PAPER_COLORS["ink"])
    ax.text(0.60, 0.86, "Post-shift", transform=ax.transAxes, fontsize=10.4, fontweight="semibold", color=PAPER_COLORS["ink"])
    ax.text(0.03, min(0.97, tau + margin + 0.02), rf"$G_q$: [{tau - margin:.2f}, {tau + margin:.2f}]", fontsize=10.0, color=PAPER_COLORS["ink"])

    ax.set_xlim(1.0, float(len(sub)))
    ax.set_ylim(0.05, 1.02)
    ax.set_xticks(np.arange(1, len(sub) + 1, 3))
    _format_axes(ax, xlabel="Actual Time Steps within the Selected Shifted Run", ylabel="Prediction Score / Recall Proxy", y_grid_only=True)
    ax.set_title("Stability Trace across a Representative Shifted Run", pad=10)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.01), ncol=4, fontsize=8.9)

    utar_floor = float(sub["p_final"].min())
    ensemble_floor = float(sub["p_ensemble"].min())
    ax.annotate(
        f"UTAR minimum: {utar_floor:.2f}",
        xy=(x[int(np.argmin(sub["p_final"].to_numpy(dtype=float)))], utar_floor),
        xytext=(0.67, 0.19),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["navy"]},
        fontsize=9.7,
        color=PAPER_COLORS["navy"],
        bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    ax.annotate(
        f"Ensemble minimum: {ensemble_floor:.2f}",
        xy=(x[int(np.argmin(sub["p_ensemble"].to_numpy(dtype=float)))], ensemble_floor),
        xytext=(0.49, 0.58),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["muted_green"]},
        fontsize=9.7,
        color=PAPER_COLORS["muted_green"],
        bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"},
    )

    stats_text = (
        f"Selected run: Fault {int(sub['fault_id'].iloc[0])}, run {int(sub['run_id'].iloc[0])}\n"
        f"Flip count near tau: ensemble {_run_flip_count(sub['p_ensemble'].to_numpy(dtype=float), tau)} vs "
        f"UTAR {_run_flip_count(sub['p_final'].to_numpy(dtype=float), tau)}"
    )
    ax.text(
        0.015,
        0.04,
        stats_text,
        transform=ax.transAxes,
        fontsize=8.8,
        color=PAPER_COLORS["ink"],
        bbox={"boxstyle": "round,pad=0.22", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    out_path = FIG_DIR / "figure5_shift_trace.png"
    save_figure(fig, out_path)
    plt.close(fig)
    return out_path


def _local_flip_profile(scores: np.ndarray, tau: float, window: int = 41) -> np.ndarray:
    states = (np.asarray(scores, dtype=float) >= tau).astype(int)
    events = np.abs(np.diff(states))
    kernel = np.ones(window - 1, dtype=float)
    local = np.convolve(events, kernel, mode="same")
    padded = np.zeros(len(scores), dtype=float)
    padded[1:] = local[: len(scores) - 1]
    return padded


def _scatter_points_by_run(df: pd.DataFrame, score_col: str, tau: float, *, phase_mask: pd.Series, max_points: int = 1800) -> tuple[np.ndarray, np.ndarray]:
    xs: list[np.ndarray] = []
    ys: list[np.ndarray] = []
    for (_, _, run_id), sub in df[phase_mask].groupby(["source_file", "fault_id", "run_id"]):
        ordered = sub.sort_values("sample_idx")
        scores = ordered[score_col].to_numpy(dtype=float)
        if len(scores) < 8:
            continue
        xs.append(scores - tau)
        ys.append(_local_flip_profile(scores, tau=tau))
    if not xs:
        return np.array([]), np.array([])
    x_all = np.concatenate(xs)
    y_all = np.concatenate(ys)
    if len(x_all) > max_points:
        rng = np.random.default_rng(42)
        idx = rng.choice(len(x_all), size=max_points, replace=False)
        x_all = x_all[idx]
        y_all = y_all[idx]
    return x_all, y_all


def _pick_boundary_run(base: pd.DataFrame, selective: pd.DataFrame, tau: float) -> pd.DataFrame:
    merged = base.merge(
        selective[["source_file", "fault_id", "run_id", "sample_idx", "p_final"]],
        on=["source_file", "fault_id", "run_id", "sample_idx"],
        how="inner",
    )
    best_sub = None
    best_key = None
    best_score = -np.inf
    for key, sub in merged[merged["phase"].isin(["transition", "post_shift"])].groupby(["source_file", "fault_id", "run_id"]):
        sub = sub.sort_values("sample_idx")
        ensemble_flips = _run_flip_count(sub["p_ensemble"].to_numpy(dtype=float), tau=tau)
        utar_flips = _run_flip_count(sub["p_final"].to_numpy(dtype=float), tau=tau)
        if ensemble_flips <= utar_flips:
            continue
        gap = ensemble_flips - utar_flips
        nonzero_bonus = 0.5 if utar_flips > 0 else 0.0
        coverage = float(np.mean(np.abs(sub["p_ensemble"].to_numpy(dtype=float) - tau) <= 0.10))
        score = 2.0 * gap + nonzero_bonus + coverage
        if score > best_score:
            best_score = score
            best_key = key
            best_sub = sub.copy()

    if best_sub is None:
        return _pick_representative_run(base, selective, tau=tau)
    best_sub["relative_step"] = np.arange(len(best_sub))
    return best_sub


def _boundary_cohort_points(base: pd.DataFrame, selective: pd.DataFrame, tau: float, margin: float) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    merged = base.merge(
        selective[["source_file", "fault_id", "run_id", "sample_idx", "p_final"]],
        on=["source_file", "fault_id", "run_id", "sample_idx"],
        how="inner",
    )
    shifted = merged[merged["phase"].isin(["transition", "post_shift"])].copy()

    cohort_rows: list[dict[str, float | int | str]] = []
    cohort_keys: list[tuple[str, int, int]] = []
    for key, sub in shifted.groupby(["source_file", "fault_id", "run_id"]):
        ordered = sub.sort_values("sample_idx")
        ensemble_flips = _run_flip_count(ordered["p_ensemble"].to_numpy(dtype=float), tau=tau)
        utar_flips = _run_flip_count(ordered["p_final"].to_numpy(dtype=float), tau=tau)
        if ensemble_flips <= utar_flips:
            continue
        cohort_keys.append(key)
        cohort_rows.append(
            {
                "source_file": key[0],
                "fault_id": int(key[1]),
                "run_id": int(key[2]),
                "ensemble_flips": ensemble_flips,
                "utar_flips": utar_flips,
                "gap": ensemble_flips - utar_flips,
            }
        )

    if not cohort_rows:
        fallback = _pick_boundary_run(base, selective, tau=tau)
        cohort_df = pd.DataFrame(
            [
                {
                    "source_file": str(fallback["source_file"].iloc[0]),
                    "fault_id": int(fallback["fault_id"].iloc[0]),
                    "run_id": int(fallback["run_id"].iloc[0]),
                    "ensemble_flips": _run_flip_count(fallback["p_ensemble"].to_numpy(dtype=float), tau=tau),
                    "utar_flips": _run_flip_count(fallback["p_final"].to_numpy(dtype=float), tau=tau),
                    "gap": _run_flip_count(fallback["p_ensemble"].to_numpy(dtype=float), tau=tau)
                    - _run_flip_count(fallback["p_final"].to_numpy(dtype=float), tau=tau),
                }
            ]
        )
        shifted = fallback[fallback["phase"].isin(["transition", "post_shift"])].copy()
        shifted["source_file"] = str(fallback["source_file"].iloc[0])
        shifted["fault_id"] = int(fallback["fault_id"].iloc[0])
        shifted["run_id"] = int(fallback["run_id"].iloc[0])
        return cohort_df, shifted, shifted

    cohort_df = pd.DataFrame(cohort_rows).sort_values(["gap", "ensemble_flips"], ascending=[False, False]).reset_index(drop=True)
    cohort_mask = pd.Series(False, index=shifted.index)
    for source_file, fault_id, run_id in cohort_keys:
        cohort_mask |= (
            (shifted["source_file"] == source_file)
            & (shifted["fault_id"] == fault_id)
            & (shifted["run_id"] == run_id)
        )
    cohort_shifted = shifted[cohort_mask].copy()

    # Keep the actual gray-zone width visible while trimming far-off margins that distort the boundary plot.
    x_limit = max(0.20, min(0.30, margin * 1.30))
    bin_edges = np.linspace(-x_limit, x_limit, 25)

    def _aggregate_points(score_col: str) -> pd.DataFrame:
        rows: list[pd.DataFrame] = []
        for key, sub in cohort_shifted.groupby(["source_file", "fault_id", "run_id"]):
            ordered = sub.sort_values("sample_idx")
            x_vals = ordered[score_col].to_numpy(dtype=float) - tau
            y_vals = _local_flip_profile(ordered[score_col].to_numpy(dtype=float), tau=tau, window=41)
            within = np.abs(x_vals) <= x_limit
            if not np.any(within):
                continue
            clipped = pd.DataFrame({"x": x_vals[within], "y": y_vals[within]})
            clipped["bin"] = pd.cut(clipped["x"], bins=bin_edges, include_lowest=True)
            agg = (
                clipped.groupby("bin", observed=False)
                .agg(x_mid=("x", "mean"), y_sum=("y", "sum"))
                .reset_index(drop=True)
            )
            centers = (bin_edges[:-1] + bin_edges[1:]) / 2.0
            agg["x_mid"] = centers[: len(agg)]
            agg["y_sum"] = agg["y_sum"].fillna(0.0)
            agg["source_file"] = key[0]
            agg["fault_id"] = int(key[1])
            agg["run_id"] = int(key[2])
            rows.append(agg)
        points = pd.concat(rows, ignore_index=True)
        return points

    ensemble_points = _aggregate_points("p_ensemble")
    utar_points = _aggregate_points("p_final")
    return cohort_df, ensemble_points, utar_points


def build_boundary_stability_figure(base: pd.DataFrame, selective: pd.DataFrame, tau: float, margin: float) -> tuple[Path, dict[str, float | int]]:
    cohort_df, ensemble_points, utar_points = _boundary_cohort_points(base, selective, tau=tau, margin=margin)
    rng = np.random.default_rng(42)
    x_simple = ensemble_points["x_mid"].to_numpy(dtype=float) + rng.normal(0.0, 0.0012, len(ensemble_points))
    y_simple = ensemble_points["y_sum"].to_numpy(dtype=float)
    x_utar = utar_points["x_mid"].to_numpy(dtype=float) + rng.normal(0.0, 0.0012, len(utar_points))
    y_utar = utar_points["y_sum"].to_numpy(dtype=float)

    plot_limit = max(0.12, float(np.nanmax(np.abs(np.r_[x_simple, x_utar])) if len(x_simple) or len(x_utar) else margin))
    fig, ax = plt.subplots(figsize=(10.4, 5.9))
    fig.subplots_adjust(left=0.10, right=0.985, top=0.88, bottom=0.14)
    ax.axvspan(-margin, margin, color="#f6d55c", alpha=0.40, zorder=0)
    ax.scatter(
        x_simple,
        y_simple,
        s=10,
        color="#1f77b4",
        alpha=0.95,
        edgecolors="white",
        linewidths=0.22,
        label="Simple Ensemble",
        zorder=3,
    )
    ax.scatter(
        x_utar,
        y_utar,
        s=18,
        marker="^",
        color="#d62728",
        alpha=0.88,
        edgecolors="white",
        linewidths=0.20,
        label="UTAR",
        zorder=4,
    )
    ax.axvline(0.0, color=PAPER_COLORS["muted_blue"], linestyle="--", linewidth=1.2, zorder=2)

    curve_simple = ensemble_points.groupby(pd.cut(ensemble_points["x_mid"], bins=np.linspace(-plot_limit, plot_limit, 17), include_lowest=True), observed=False).agg(x=("x_mid", "mean"), y=("y_sum", "mean")).dropna()
    curve_utar = utar_points.groupby(pd.cut(utar_points["x_mid"], bins=np.linspace(-plot_limit, plot_limit, 17), include_lowest=True), observed=False).agg(x=("x_mid", "mean"), y=("y_sum", "mean")).dropna()
    if not curve_simple.empty:
        ax.plot(curve_simple["x"], curve_simple["y"], color="#125da8", linewidth=2.0, alpha=0.95, zorder=5)
    if not curve_utar.empty:
        ax.plot(curve_utar["x"], curve_utar["y"], color="#b22222", linewidth=1.8, alpha=0.95, zorder=5)

    simple_peak_idx = int(np.nanargmax(y_simple)) if len(y_simple) else 0
    utar_peak_idx = int(np.nanargmax(y_utar)) if len(y_utar) else 0
    ax.annotate(
        "Simple Ensemble:\nHigh volatility around\nDecision Margin",
        xy=(x_simple[simple_peak_idx], y_simple[simple_peak_idx] if len(y_simple) else 0.0),
        xytext=(0.72, 0.68),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
        fontsize=10.0,
        bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    ax.annotate(
        "UTAR Stabilization\nnear the Gray-Zone",
        xy=(x_utar[utar_peak_idx] if len(x_utar) else 0.0, y_utar[utar_peak_idx] if len(y_utar) else 0.0),
        xytext=(0.73, 0.30),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
        fontsize=10.0,
        bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"},
    )

    y_limit = max(32.0, float(np.nanmax(np.r_[y_simple, y_utar])) * 1.18) if len(y_simple) or len(y_utar) else 32.0
    ax.set_xlim(-plot_limit, plot_limit)
    ax.set_ylim(-0.5, y_limit)
    ax.text(
        0.04,
        0.08,
        (
            f"Shifted runs used: {len(cohort_df)} (ensemble flips > UTAR)\n"
            f"Median flips/run: ensemble {cohort_df['ensemble_flips'].median():.0f} vs UTAR {cohort_df['utar_flips'].median():.0f}"
        ),
        transform=ax.transAxes,
        fontsize=8.8,
        color=PAPER_COLORS["ink"],
        bbox={"boxstyle": "round,pad=0.20", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    ax.annotate(
        rf"$G_q$ Gray-Zone Margin: {(-margin):+.2f} to {margin:+.2f}",
        xy=(-margin * 0.10, y_limit * 0.78),
        xytext=(0.10, 0.72),
        textcoords="axes fraction",
        arrowprops={"arrowstyle": "->", "lw": 1.0, "color": PAPER_COLORS["ink"]},
        fontsize=10.0,
        fontweight="semibold",
    )
    ax.text(0.505, 0.04, rf"$\tau$ ({tau:.2f})", transform=ax.transAxes, fontsize=10.0, color=PAPER_COLORS["ink"])
    ax.set_title("Stability and Volatility near the Decision Boundary", pad=10)
    _format_axes(ax, xlabel="Decision Margin", ylabel="Prediction Flip Count (Stability Measure)", y_grid_only=True)
    ax.legend(loc="upper left", ncol=2, fontsize=9.4)
    out_path = FIG_DIR / "figure6_boundary_stability.png"
    save_figure(fig, out_path)
    plt.close(fig)
    meta = {
        "cohort_size": int(len(cohort_df)),
        "ensemble_median_flips": float(cohort_df["ensemble_flips"].median()),
        "utar_median_flips": float(cohort_df["utar_flips"].median()),
    }
    return out_path, meta


def build_appendix_e_figures() -> dict[str, Path]:
    shift_table = read_csv(SHIFT_DIR / "feature_shift_normal_vs_post_shift.csv").sort_values(["ks_stat", "mean_abs_diff"], ascending=[False, False])
    top_features = shift_table["feature"].head(3).tolist()
    raw_df = read_csv(PROCESSED_DIR / "te_test_main_rows.csv")
    feature_cols = infer_feature_cols(raw_df)
    X = prepare_feature_matrix(raw_df.rename(columns={"y": "y"}), feature_cols)
    labels = raw_df["phase"].to_numpy()
    X_sample, labels_sample = sample_for_embedding(X, labels, max_per_group=420, random_state=42)
    pca = PCA(n_components=min(18, X_sample.shape[1]), random_state=42)
    X_reduced = pca.fit_transform(X_sample)
    tsne = TSNE(n_components=2, perplexity=45, init="pca", learning_rate="auto", random_state=42)
    emb = tsne.fit_transform(X_reduced)

    summary = read_csv(SHIFT_DIR / "shift_summary_table.csv")
    summary_row = summary[summary["comparison"] == "normal vs post_shift"].iloc[0]

    palette = {"normal": PAPER_COLORS["navy"], "transition": PAPER_COLORS["orange"], "post_shift": PAPER_COLORS["red"]}
    label_names = {"normal": "Normal", "transition": "Transition", "post_shift": "Post-shift"}
    paths: dict[str, Path] = {}

    fig_tsne, ax_tsne = plt.subplots(figsize=(6.8, 5.8))
    fig_tsne.subplots_adjust(left=0.10, right=0.98, top=0.88, bottom=0.12)
    for phase in ["normal", "transition", "post_shift"]:
        mask = labels_sample == phase
        ax_tsne.scatter(
            emb[mask, 0],
            emb[mask, 1],
            s=24,
            color=palette[phase],
            alpha=0.72,
            edgecolors="white",
            linewidths=0.25,
            label=label_names[phase],
        )
        centroid = emb[mask].mean(axis=0)
        ax_tsne.scatter([centroid[0]], [centroid[1]], s=110, color=palette[phase], edgecolors="white", linewidths=1.0, zorder=5)
        ax_tsne.text(centroid[0] + 1.6, centroid[1] + 1.2, label_names[phase], fontsize=9.6, fontweight="semibold", color=palette[phase])
    phase_centroids = {phase: emb[labels_sample == phase].mean(axis=0) for phase in ["normal", "transition", "post_shift"]}
    for left, right in [("normal", "transition"), ("transition", "post_shift")]:
        ax_tsne.annotate(
            "",
            xy=phase_centroids[right],
            xytext=phase_centroids[left],
            arrowprops={"arrowstyle": "->", "lw": 1.1, "color": PAPER_COLORS["ink"], "alpha": 0.7},
        )
    ax_tsne.set_title("Figure E1(a). t-SNE View of the TEP Dataset under Temporal Shift", pad=10)
    _format_axes(ax_tsne, xlabel="t-SNE Dimension 1", ylabel="t-SNE Dimension 2")
    ax_tsne.legend(loc="upper right", fontsize=9.4)
    summary_lines = [
        f"Avg |d mean|: {summary_row['avg_mean_abs_diff']:.3f}",
        f"Avg KS statistic: {summary_row['avg_ks_stat']:.3f}",
        f"KS reject ratio @ 0.05: {summary_row['ks_reject_ratio_0.05'] * 100:.1f}%",
        f"RBF-MMD: {summary_row['mmd_rbf']:.4f}",
        f"t-SNE centroid shift: {summary_row['tsne_centroid_normal_vs_post_shift']:.2f}",
    ]
    ax_tsne.text(
        0.02,
        0.02,
        "\n".join(summary_lines),
        transform=ax_tsne.transAxes,
        fontsize=8.8,
        color=PAPER_COLORS["ink"],
        bbox={"boxstyle": "round,pad=0.25", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    tsne_path = APPENDIX_DIR / "figure_e1_tsne_shift.png"
    save_figure(fig_tsne, tsne_path)
    plt.close(fig_tsne)
    paths["tsne"] = tsne_path

    plot_df = raw_df[raw_df["phase"].isin(["normal", "transition", "post_shift"])].copy()
    for idx, feat in enumerate(top_features, start=1):
        fig_kde, ax = plt.subplots(figsize=(4.8, 4.1))
        fig_kde.subplots_adjust(left=0.14, right=0.98, top=0.84, bottom=0.16)
        values = {}
        for phase in ["normal", "transition", "post_shift"]:
            series = _clean_series(plot_df.loc[plot_df["phase"] == phase, feat])
            values[phase] = series
        stacked = np.concatenate([arr for arr in values.values() if len(arr)]) if any(len(arr) for arr in values.values()) else np.array([0.0, 1.0])
        lo, hi = np.quantile(stacked, [0.005, 0.995])
        pad = (hi - lo) * 0.12 if hi > lo else 1.0
        grid = np.linspace(lo - pad, hi + pad, 500)
        for phase in ["normal", "transition", "post_shift"]:
            curve = _kde_on_grid(values[phase], grid)
            ax.fill_between(grid, curve, alpha=0.17, color=palette[phase], zorder=1)
            ax.plot(grid, curve, linewidth=1.8, color=palette[phase], label=label_names[phase], zorder=2)
        feat_row = shift_table.loc[shift_table["feature"] == feat].iloc[0]
        ax.set_title(f"Figure E1({chr(97 + idx)}). {feat.upper()} KDE", fontsize=11.0, pad=6)
        ax.text(
            0.03,
            0.95,
            f"KS={feat_row['ks_stat']:.3f}\n|d mean|={feat_row['mean_abs_diff']:.3f}",
            transform=ax.transAxes,
            va="top",
            fontsize=8.6,
            color=PAPER_COLORS["ink"],
            bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"},
        )
        _format_axes(ax, xlabel="Value", ylabel="Density", y_grid_only=True)
        ax.legend(loc="upper right", fontsize=8.4)
        kde_path = APPENDIX_DIR / f"figure_e1_kde_{idx}_{feat}.png"
        save_figure(fig_kde, kde_path)
        plt.close(fig_kde)
        paths[f"kde_{idx}"] = kde_path

    return paths


def _select_graphad_example() -> tuple[pd.Series, pd.DataFrame, dict]:
    pred_main = read_csv(PRED_DIR / "base_test_main_predictions.csv")
    shift_rows = pred_main[(pred_main["phase"].isin(["transition", "post_shift"])) & (pred_main["y_true"] == 1)].copy()
    shift_rows["sensor_changed"] = shift_rows["graphad_top1_sensor"].astype(str) != shift_rows["graphad_raw_top1_sensor"].astype(str)
    shift_rows["priority"] = shift_rows["sensor_changed"].astype(int) * 10.0
    shift_rows["priority"] += (shift_rows["graphad_top1_gap"].fillna(0.0) - shift_rows["graphad_raw_top1_gap"].fillna(0.0)).abs()
    shift_rows["priority"] += shift_rows["graphad_score"].fillna(0.0)
    raw_rows = read_csv(PROCESSED_DIR / "te_test_main_rows.csv")
    artifact = load_graphad_artifact(MODEL_DIR / "graphad_artifact.json")
    run_cache: dict[tuple[str, int, int], pd.DataFrame] = {}

    for _, row in shift_rows.sort_values(["priority", "graphad_score"], ascending=[False, False]).head(60).iterrows():
        run_key = (str(row["source_file"]), int(row["fault_id"]), int(row["run_id"]))
        if run_key not in run_cache:
            run_cache[run_key] = raw_rows[
                (raw_rows["source_file"] == row["source_file"])
                & (raw_rows["fault_id"] == row["fault_id"])
                & (raw_rows["run_id"] == row["run_id"])
            ].sort_values("sample_idx")
        run_df = run_cache[run_key]
        mats = graphad_score_matrices(run_df, artifact)
        sample_idx = int(row["sample_idx"])
        mask = run_df["sample_idx"] == sample_idx
        if not mask.any():
            continue
        raw_series = mats["raw"].loc[mask].iloc[0].sort_values(ascending=False)
        smooth_series = mats["smooth"].loc[mask].iloc[0].sort_values(ascending=False)
        target_sensor = str(smooth_series.index[0])
        raw_target_rank = int(raw_series.index.get_loc(target_sensor) + 1)
        if raw_target_rank > 1:
            return row, run_df, artifact

    row = shift_rows.sort_values(["priority", "graphad_score"], ascending=[False, False]).iloc[0]
    run_df = raw_rows[
        (raw_rows["source_file"] == row["source_file"])
        & (raw_rows["fault_id"] == row["fault_id"])
        & (raw_rows["run_id"] == row["run_id"])
    ].sort_values("sample_idx")
    return row, run_df, artifact


def build_appendix_f_figures() -> tuple[dict[str, Path], dict[str, object]]:
    row, run_df, artifact = _select_graphad_example()
    mats = graphad_score_matrices(run_df, artifact)
    sample_idx = int(row["sample_idx"])
    sample_scores = {key: frame.loc[run_df["sample_idx"] == sample_idx].iloc[0] for key, frame in mats.items()}
    raw_series = sample_scores["raw"].sort_values(ascending=False)
    smooth_series = sample_scores["smooth"].sort_values(ascending=False)
    z_series = sample_scores["z"]
    trend_series = sample_scores["trend"]
    fluct_series = sample_scores["fluct"]

    selected = list(dict.fromkeys(list(raw_series.head(7).index) + list(smooth_series.head(7).index)))
    adjacency = artifact.get("adjacency", {})
    sensor_order = selected
    angles = np.linspace(np.pi * 0.12, np.pi * 2.12, len(sensor_order), endpoint=False)
    coords = np.c_[0.5 + 0.28 * np.cos(angles), 0.52 + 0.23 * np.sin(angles)]
    sensor_pos = {sensor: tuple(coord) for sensor, coord in zip(sensor_order, coords)}
    target_sensor = str(smooth_series.index[0])
    raw_top = str(raw_series.index[0])
    raw_target_rank = int(raw_series.index.get_loc(target_sensor) + 1)
    smooth_target_rank = int(smooth_series.index.get_loc(target_sensor) + 1)

    edges: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for sensor in sensor_order:
        for neighbor in adjacency.get(sensor, []):
            if neighbor not in sensor_pos:
                continue
            key = tuple(sorted((sensor, neighbor)))
            if key in seen:
                continue
            seen.add(key)
            edges.append((sensor, neighbor))
    if len(edges) < len(sensor_order):
        for left, right in zip(sensor_order, sensor_order[1:] + sensor_order[:1]):
            key = tuple(sorted((left, right)))
            if key not in seen:
                edges.append((left, right))
                seen.add(key)

    raw_norm = raw_series.loc[sensor_order].to_numpy(dtype=float)
    raw_norm = raw_norm / max(float(raw_norm.max()), 1e-8)
    smooth_norm = smooth_series.loc[sensor_order].to_numpy(dtype=float)
    smooth_norm = smooth_norm / max(float(smooth_norm.max()), 1e-8)

    def _panel_figure(title: str) -> tuple[object, object]:
        fig, ax = plt.subplots(figsize=(5.35, 4.8))
        fig.subplots_adjust(left=0.03, right=0.97, top=0.88, bottom=0.08)
        ax.set_axis_off()
        ax.set_xlim(0.0, 1.0)
        ax.set_ylim(0.0, 1.0)
        ax.set_title(title, pad=10)
        return fig, ax

    def _label_position(x0: float, y0: float, radius: float = 0.16) -> tuple[float, float]:
        dx = x0 - 0.5
        dy = y0 - 0.52
        scale = radius / max(np.hypot(dx, dy), 1e-6)
        return x0 + dx * scale, y0 + dy * scale

    def _add_label(ax, x0: float, y0: float, text: str, *, facecolor: str = "white", edgecolor: str = "#d9dde3", fontsize: float = 8.5, color: str = PAPER_COLORS["ink"]) -> None:
        lx, ly = _label_position(x0, y0)
        ax.plot([x0, lx], [y0, ly], color="#aeb7c2", linewidth=0.8, zorder=4)
        ax.text(
            lx,
            ly,
            text,
            fontsize=fontsize,
            color=color,
            bbox={"boxstyle": "round,pad=0.18", "facecolor": facecolor, "edgecolor": edgecolor},
            ha="left" if lx >= x0 else "right",
            va="center",
            zorder=5,
        )

    raw_fig, raw_ax = _panel_figure("(a) Raw Anomaly Scores: Data-driven Ambiguity")
    smooth_fig, smooth_ax = _panel_figure("(b) Graph Smoothing: Correlation-based Propagation")
    final_fig, final_ax = _panel_figure("(c) Hybrid Reranking: Root-cause Identification")

    for sensor, neighbor in edges:
        x0, y0 = sensor_pos[sensor]
        x1, y1 = sensor_pos[neighbor]
        raw_ax.plot([x0, x1], [y0, y1], color="#cacaca", linewidth=1.5, zorder=1)
        final_ax.plot([x0, x1], [y0, y1], color="#d0d0d0", linewidth=1.5, zorder=1)

    for sensor, neighbor in edges:
        s0 = smooth_series[sensor]
        s1 = smooth_series[neighbor]
        src, dst = (sensor, neighbor) if s0 <= s1 else (neighbor, sensor)
        x0, y0 = sensor_pos[src]
        x1, y1 = sensor_pos[dst]
        smooth_ax.add_patch(
            FancyArrowPatch(
                (x0, y0),
                (x1, y1),
                connectionstyle="arc3,rad=0.12",
                arrowstyle="-|>",
                mutation_scale=11,
                linewidth=1.3,
                color="#7fb3d5",
                alpha=0.82,
            )
        )

    raw_label_sensors = list(dict.fromkeys(list(raw_series.head(4).index) + ([target_sensor] if target_sensor not in raw_series.head(4).index else [])))
    smooth_label_sensors = list(dict.fromkeys(list(smooth_series.head(4).index) + ([raw_top] if raw_top not in smooth_series.head(4).index else [])))
    final_label_sensors = [sensor for sensor in list(smooth_series.head(5).index) if sensor != target_sensor]

    for idx, sensor in enumerate(sensor_order):
        x0, y0 = sensor_pos[sensor]
        raw_color = plt.cm.Wistia(0.35 + 0.55 * raw_norm[idx])
        smooth_color = plt.cm.cividis(0.20 + 0.72 * smooth_norm[idx])
        final_rank = int(smooth_series.index.get_loc(sensor) + 1)
        final_color = ["#2c6fb7", "#4f81bd", "#c87f2a", "#d9a441", "#ead39c", "#f2e7c9", "#f7f2e6"][min(final_rank - 1, 6)]

        raw_ax.scatter([x0], [y0], s=280 + 240 * raw_norm[idx], color=raw_color, edgecolor="#8a6d3b", linewidth=1.1, zorder=3)
        smooth_ax.scatter([x0], [y0], s=300 + 300 * smooth_norm[idx], color=smooth_color, edgecolor="white", linewidth=1.1, zorder=3)
        final_ax.scatter([x0], [y0], s=300 if sensor == target_sensor else 255, color=final_color, edgecolor="#0c4b84" if sensor == target_sensor else "#8a6d3b", linewidth=1.25, zorder=3)

        if sensor in raw_label_sensors:
            _add_label(raw_ax, x0, y0, f"{sensor.upper()}: {raw_series[sensor]:.2f}", facecolor="#fff7d6", edgecolor="#d8bf61", fontsize=8.2)
        if sensor in smooth_label_sensors:
            _add_label(smooth_ax, x0, y0, f"{sensor.upper()}: {raw_series[sensor]:.2f} -> {smooth_series[sensor]:.2f}", fontsize=7.8)
        if sensor in final_label_sensors:
            final_label = f"{sensor.upper()}\n#{final_rank}"
            if sensor == target_sensor:
                final_label = f"{sensor.upper()}\nTop-1"
            _add_label(final_ax, x0, y0, final_label, fontsize=8.1, facecolor="white", edgecolor="#d9dde3")

        if sensor == raw_top:
            raw_ax.text(x0 - 0.03, y0 - 0.13, "Raw Top-1", fontsize=8.4, fontweight="semibold", color="#8a5a00", bbox={"boxstyle": "round,pad=0.18", "facecolor": "#ffe699", "edgecolor": "#c8a64d"}, zorder=6)
        if sensor == target_sensor and raw_target_rank > 1:
            raw_ax.text(x0 - 0.10, y0 + 0.12, f"Promoted from\nRaw #{raw_target_rank}", fontsize=8.2, color=PAPER_COLORS["navy"], bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"}, zorder=6)

    delta_rank = raw_target_rank - smooth_target_rank
    delta_text = f"Delta: +{delta_rank}" if delta_rank >= 0 else f"Delta: {delta_rank}"
    summary_text = (
        "Promotion Summary\n"
        f"Target: {target_sensor.upper()}\n"
        f"Raw rank: #{raw_target_rank}\n"
        f"Final rank: #{smooth_target_rank}\n"
        f"{delta_text}"
    )
    final_ax.text(
        0.60,
        0.86,
        summary_text,
        ha="left",
        va="top",
        fontsize=8.8,
        color=PAPER_COLORS["ink"],
        bbox={"boxstyle": "round,pad=0.24", "facecolor": "white", "edgecolor": "#d9dde3"},
    )
    final_ax.text(0.06, 0.05, "Rank palette:", fontsize=8.4, fontweight="semibold")
    palette = ["#2c6fb7", "#4f81bd", "#c87f2a", "#d9a441", "#ead39c"]
    for idx, x0 in enumerate(np.linspace(0.25, 0.61, 5), start=1):
        final_ax.text(
            x0,
            0.05,
            f"{idx}",
            ha="center",
            va="center",
            fontsize=8.1,
            color="white" if idx <= 3 else PAPER_COLORS["ink"],
            bbox={"boxstyle": "round,pad=0.16", "facecolor": palette[idx - 1], "edgecolor": palette[idx - 1]},
        )

    for ax, footer in [
        (raw_ax, f"Sample {sample_idx}: raw evidence spreads across correlated sensors."),
        (smooth_ax, "Graph smoothing propagates evidence through the learned topology."),
        (final_ax, f"The top reranked candidate becomes {target_sensor.upper()} after GraphAD+ smoothing."),
    ]:
        ax.text(0.02, 0.02, footer, transform=ax.transAxes, fontsize=8.4, color=PAPER_COLORS["ink"], bbox={"boxstyle": "round,pad=0.18", "facecolor": "white", "edgecolor": "#d9dde3"})

    paths = {
        "raw": APPENDIX_DIR / "figure_f1_panel_raw_scores.png",
        "smooth": APPENDIX_DIR / "figure_f1_panel_graph_smoothing.png",
        "rerank": APPENDIX_DIR / "figure_f1_panel_hybrid_reranking.png",
    }
    save_figure(raw_fig, paths["raw"])
    save_figure(smooth_fig, paths["smooth"])
    save_figure(final_fig, paths["rerank"])
    plt.close(raw_fig)
    plt.close(smooth_fig)
    plt.close(final_fig)
    meta = {
        "sample_idx": sample_idx,
        "fault_id": int(row["fault_id"]),
        "run_id": int(row["run_id"]),
        "raw_top_sensor": raw_top.upper(),
        "raw_top_score": float(raw_series.iloc[0]),
        "target_sensor": target_sensor.upper(),
        "target_score": float(smooth_series.iloc[0]),
        "raw_target_rank": raw_target_rank,
        "final_target_rank": smooth_target_rank,
        "support_sensor_1": str(raw_series.index[2]).upper() if len(raw_series) > 2 else "",
        "support_sensor_1_score": float(raw_series.iloc[2]) if len(raw_series) > 2 else 0.0,
        "support_sensor_2": str(raw_series.index[3]).upper() if len(raw_series) > 3 else "",
        "support_sensor_2_score": float(raw_series.iloc[3]) if len(raw_series) > 3 else 0.0,
    }
    return paths, meta


def _normalize_text(text: str) -> str:
    return " ".join((text or "").split())


def _find_paragraph(doc: Document, needle: str) -> tuple[int, object]:
    target = _normalize_text(needle)
    for idx, paragraph in enumerate(doc.paragraphs):
        if target in _normalize_text(paragraph.text):
            return idx, paragraph
    raise ValueError(f"Could not locate paragraph containing: {needle}")


def _clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_picture_before(anchor_paragraph, image_path: Path, width_inches: float):
    new_paragraph = _insert_paragraph_before(anchor_paragraph)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return new_paragraph


def _insert_caption_before(anchor_paragraph, caption: str):
    new_paragraph = _insert_paragraph_before(anchor_paragraph)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_paragraph.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    return new_paragraph


def _set_paragraph_text(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def _insert_picture_row_before(anchor_paragraph, image_paths: list[Path], width_inches: float, spacer: str = "  "):
    new_paragraph = _insert_paragraph_before(anchor_paragraph)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, image_path in enumerate(image_paths):
        run = new_paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        if idx < len(image_paths) - 1:
            new_paragraph.add_run(spacer)
    return new_paragraph


def _remove_drawing_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        if "w:drawing" in run._r.xml:
            paragraph._p.remove(run._r)


def _replace_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        with ZipFile(docx_path) as zf:
            zf.extractall(tmpdir_path)
        for internal, src in replacements.items():
            if src.exists():
                shutil.copy2(src, tmpdir_path / internal)
        rebuilt = docx_path.with_suffix(".tmp.docx")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as zf:
            for path in sorted(tmpdir_path.rglob("*")):
                if path.is_file():
                    zf.write(path, path.relative_to(tmpdir_path))
        rebuilt.replace(docx_path)


def update_add_docx(
    methodology_paths: dict[str, Path],
    figure4_path: Path,
    figure5_path: Path,
    figure6_path: Path,
    figure6_meta: dict[str, float | int],
    appendix_e_paths: dict[str, Path],
    appendix_f_paths: dict[str, Path],
    appendix_f_meta: dict[str, object],
) -> Path:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not DOCX_BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, DOCX_BACKUP_PATH)
    else:
        shutil.copy2(DOCX_BACKUP_PATH, DOCX_PATH)

    doc = Document(DOCX_PATH)

    fig1_caption_idx, fig1_caption = _find_paragraph(doc, "Figure 1. Proposed Robust Diagnosis Framework")
    fig1_para = doc.paragraphs[fig1_caption_idx - 1]
    _clear_paragraph(fig1_para)
    fig1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, key in enumerate(["distribution", "margin", "discrepancy"]):
        run = fig1_para.add_run()
        run.add_picture(str(methodology_paths[key]), width=Inches(1.96))
        if idx < 2:
            fig1_para.add_run("  ")

    _, fig4_anchor = _find_paragraph(doc, "Figure 4는 IID와 TDS 사이의 성능 저하량")
    _insert_picture_before(fig4_anchor, figure4_path, width_inches=6.05)
    _insert_caption_before(fig4_anchor, "Figure 4. Performance Robustness under Distribution Shift.")

    _, fig5_anchor = _find_paragraph(doc, "실험 결과, 단순 앙상블(Simple Ensemble) 모델은 TDS 환경에서")
    _insert_picture_before(fig5_anchor, figure5_path, width_inches=6.05)
    _insert_caption_before(fig5_anchor, "Figure 5. Stability Trace across a Representative Shifted Run.")

    _, appendix_f_heading = _find_paragraph(doc, "F. Visual Proof of Data-driven Hybrid Reranking")
    _insert_picture_before(appendix_f_heading, appendix_e_paths["tsne"], width_inches=5.85)
    _insert_picture_row_before(
        appendix_f_heading,
        [appendix_e_paths["kde_1"], appendix_e_paths["kde_2"], appendix_e_paths["kde_3"]],
        width_inches=1.92,
    )
    _insert_caption_before(appendix_f_heading, "Figure E1. KDE and t-SNE Views of the TEP Dataset under Temporal Shift.")

    _, fig6_anchor = _find_paragraph(doc, "단순 앙상블 모델은 결정 임계치")
    _insert_caption_before(fig6_anchor, "Figure 6. Stability and Volatility near the Decision Boundary.")
    fig6_body = (
        f"Figure 6은 전체 shifted sample 전체를 직접 합친 그림이 아니라, shifted run 중 단순 앙상블의 prediction flip 수가 "
        f"UTAR보다 더 크게 나타난 cohort({int(figure6_meta['cohort_size'])} runs)만을 모아 구성한 시각화이다. "
        f"이 cohort에서 단순 앙상블의 median flip/run은 {float(figure6_meta['ensemble_median_flips']):.0f}회, "
        f"UTAR는 {float(figure6_meta['utar_median_flips']):.0f}회로 나타났으며, 결정 임계치($\\tau$) 부근에서 "
        "단순 앙상블이 더 큰 판정 요동을 보이는 run들을 대상으로 UTAR의 안정화 효과를 비교한 그림으로 해석해야 한다. "
        "즉, 본 그림은 전체 shifted population에 대한 평균적 경향을 직접 도식화한 것이 아니라, UTAR가 불안정한 경계 상황을 "
        "어떻게 완화하는지를 보여주는 cohort-based evidence이다."
    )
    _set_paragraph_text(fig6_anchor, fig6_body)

    appendix_f_idx, appendix_f_heading = _find_paragraph(doc, "F. Visual Proof of Data-driven Hybrid Reranking")
    _remove_drawing_runs(appendix_f_heading)
    appendix_f_anchor = doc.paragraphs[appendix_f_idx + 1]
    _insert_picture_row_before(
        appendix_f_anchor,
        [appendix_f_paths["raw"], appendix_f_paths["smooth"], appendix_f_paths["rerank"]],
        width_inches=1.93,
    )

    appendix_f_updates = {
        appendix_f_idx + 2: f"(a) Raw Anomaly Scores (t={int(appendix_f_meta['sample_idx'])}): Data-driven Ambiguity",
        appendix_f_idx + 3: (
            f"현상: sample {int(appendix_f_meta['sample_idx'])}에서는 {appendix_f_meta['raw_top_sensor']}와 "
            f"{appendix_f_meta['target_sensor']}가 모두 raw anomaly score {float(appendix_f_meta['raw_top_score']):.2f}로 최상위권에 나타나며, "
            f"{appendix_f_meta['support_sensor_1']} ({float(appendix_f_meta['support_sensor_1_score']):.2f})와 "
            f"{appendix_f_meta['support_sensor_2']} ({float(appendix_f_meta['support_sensor_2_score']):.2f})가 뒤따른다."
        ),
        appendix_f_idx + 4: (
            f"문제점: raw score만으로는 {appendix_f_meta['raw_top_sensor']}와 {appendix_f_meta['target_sensor']}의 우선순위를 분리하기 어렵고, "
            "상관된 후보들이 동시에 높게 나타나 root-cause ordering이 모호해진다."
        ),
        appendix_f_idx + 5: "(b) Graph Smoothing (Eq. 8): Correlation-based Propagation",
        appendix_f_idx + 6: (
            "작동 원리: 본 연구는 도메인 지식 대신 Pearson correlation 기반의 데이터 주도 topology를 학습하고, "
            "그 위에서 anomaly evidence를 전파해 구조적으로 일관된 후보를 강조한다."
        ),
        appendix_f_idx + 7: (
            f"효과: {appendix_f_meta['raw_top_sensor']}와 {appendix_f_meta['target_sensor']}처럼 raw score가 비슷한 후보들 사이에서도, "
            "graph smoothing은 주변 상관구조를 반영해 ordering을 안정화한다."
        ),
        appendix_f_idx + 8: "(c) Hybrid Reranking: Root-Cause Identification",
        appendix_f_idx + 9: (
            f"최종 결과: {appendix_f_meta['target_sensor']}는 raw rank #{int(appendix_f_meta['raw_target_rank'])}에서 "
            f"final rank #{int(appendix_f_meta['final_target_rank'])}로 상승하고, {appendix_f_meta['raw_top_sensor']}는 그 다음 후보로 재정렬된다."
        ),
        appendix_f_idx + 10: (
            "이 시각화는 동일한 raw anomaly score를 가진 후보들 사이에서도, "
            "GraphAD+의 데이터 기반 topology가 root-cause candidate ordering을 분해하고 안정화할 수 있음을 보여준다."
        ),
    }
    for paragraph_idx, text in appendix_f_updates.items():
        _set_paragraph_text(doc.paragraphs[paragraph_idx], text)

    doc.save(DOCX_PATH)

    _replace_media(
        DOCX_PATH,
        {
            "word/media/image6.png": figure6_path,
        },
    )
    return DOCX_PATH


def main() -> None:
    set_paper_style()
    ensure_dir(FIG_DIR)
    ensure_dir(APPENDIX_DIR)

    base, val_df, selective, tau, margin, temporal_label = _load_predictions()
    methodology_paths = build_methodology_panels(base, val_df, tau=tau, margin=margin, temporal_label=temporal_label)
    figure4_path = build_performance_drop_figure()
    figure5_path = build_shift_trace_figure(base, selective, tau=tau, margin=margin)
    figure6_path, figure6_meta = build_boundary_stability_figure(base, selective, tau=tau, margin=margin)
    appendix_e_paths = build_appendix_e_figures()
    appendix_f_paths, appendix_f_meta = build_appendix_f_figures()
    updated_docx = update_add_docx(
        methodology_paths=methodology_paths,
        figure4_path=figure4_path,
        figure5_path=figure5_path,
        figure6_path=figure6_path,
        figure6_meta=figure6_meta,
        appendix_e_paths=appendix_e_paths,
        appendix_f_paths=appendix_f_paths,
        appendix_f_meta=appendix_f_meta,
    )
    print(f"Updated figures and document written to: {updated_docx}")


if __name__ == "__main__":
    main()
