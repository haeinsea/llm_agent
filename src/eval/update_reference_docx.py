from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile

import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from src.routing.selective_llm_eval import read_selected_q, DEFAULT_Q
from src.utils.experiment import get_seed_count
from src.utils.io import read_csv, read_json


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = PROJECT_ROOT / "outputs"
METRIC_DIR = OUTPUT_DIR / "metrics"
APPENDIX_DIR = OUTPUT_DIR / "appendix"
FIG_DIR = OUTPUT_DIR / "figures"
EVAL_DIR = OUTPUT_DIR / "evaluation"

SOURCE_DOCX = PROJECT_ROOT / "UTAR 논문(Reference)_260321.docx"
TARGET_DOCX = PROJECT_ROOT / "UTAR 논문(Reference)_260321_add.docx"


def _set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"


def _remove_row_trailing_omission(row) -> None:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return
    for tag in ("w:gridAfter", "w:wAfter"):
        el = tr_pr.find(qn(tag))
        if el is not None:
            tr_pr.remove(el)


def _append_blank_cell(row) -> None:
    _remove_row_trailing_omission(row)
    new_tc = deepcopy(row._tr.tc_lst[-1])
    for node in new_tc.iter():
        if node.tag == qn("w:t"):
            node.text = ""
    row._tr.append(new_tc)


def _split_last_header_span(row) -> None:
    if not row._tr.tc_lst:
        return
    last_tc = row._tr.tc_lst[-1]
    tc_pr = last_tc.tcPr
    if tc_pr is None:
        return
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return
    span = int(grid_span.get(qn("w:val"), "1"))
    if span <= 1:
        return
    tc_w = tc_pr.find(qn("w:tcW"))
    split_width = None
    if tc_w is not None and tc_w.get(qn("w:w")):
        try:
            width = int(tc_w.get(qn("w:w")))
            split_width = max(width // span, 1)
            tc_w.set(qn("w:w"), str(split_width))
        except ValueError:
            split_width = None
    tc_pr.remove(grid_span)
    new_tc = deepcopy(last_tc)
    new_tc_pr = new_tc.find(qn("w:tcPr"))
    if new_tc_pr is not None:
        new_grid_span = new_tc_pr.find(qn("w:gridSpan"))
        if new_grid_span is not None:
            new_tc_pr.remove(new_grid_span)
        new_tc_w = new_tc_pr.find(qn("w:tcW"))
        if new_tc_w is not None and split_width is not None:
            new_tc_w.set(qn("w:w"), str(split_width))
    row._tr.append(new_tc)


def _resize_table(table, n_rows: int, n_cols: int) -> None:
    while len(table.rows[0]._tr.tc_lst) < n_cols and len(table.rows[0].cells) >= n_cols:
        _split_last_header_span(table.rows[0])
    while len(table.rows) < n_rows:
        table.add_row()
    while len(table.rows) > n_rows:
        table._tbl.remove(table.rows[-1]._tr)
    for row in table.rows:
        while len(row.cells) < n_cols:
            _append_blank_cell(row)


def _write_dataframe_to_table(table, df: pd.DataFrame) -> None:
    text_df = df.fillna("").astype(str)
    _resize_table(table, len(text_df) + 1, len(text_df.columns))
    if len(table.rows[0].cells) < len(text_df.columns):
        raise ValueError(f"Header row has only {len(table.rows[0].cells)} cells; cannot fit {len(text_df.columns)} columns.")
    for j, col in enumerate(text_df.columns):
        _set_cell_text(table.rows[0].cells[j], col)
    for i, (_, row) in enumerate(text_df.iterrows(), start=1):
        if len(table.rows[i].cells) < len(text_df.columns):
            raise ValueError(f"Row {i} has only {len(table.rows[i].cells)} cells; cannot fit {len(text_df.columns)} columns.")
        for j, value in enumerate(row.tolist()):
            _set_cell_text(table.rows[i].cells[j], value)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _selected_summary_row(dataset: str, mode: str, q: float) -> pd.Series:
    summary = read_csv(METRIC_DIR / "selective_llm_summary.csv")
    sub = summary[(summary["dataset"] == dataset) & (summary["mode"] == mode) & (np.isclose(summary["q"], q))]
    if sub.empty:
        raise KeyError(f"No summary row for dataset={dataset}, mode={mode}, q={q}")
    return sub.iloc[0]


def _build_table1() -> pd.DataFrame:
    return read_csv(METRIC_DIR / "table1_model_configuration.csv")


def _build_table2(selected_q: float) -> pd.DataFrame:
    base = read_csv(METRIC_DIR / "table2_robustness.csv").copy()
    return base[["Method", "Ave. F1", "PRR", "Worst-Case Recall (P5)", "Instability", "Inference Time (s)", "Inference Time / Sample (ms)"]]


def _build_table3() -> pd.DataFrame:
    df = read_csv(METRIC_DIR / "table3_q_sweep.csv").copy()
    return df[["Strategy", "Call Rate", "F1-Score", "Cost (USD)", "Inference Time (s)"]]


def _build_table5(selected_q: float) -> pd.DataFrame:
    df = read_csv(METRIC_DIR / "table5_ablation.csv").copy()
    return df[["Configuration", "Ave. F1", "PRR", "Worst-Case Recall (P5)", "Instability", "Inference Time (s)", "Inference Time / Sample (ms)"]]


def _build_table6() -> pd.DataFrame:
    return read_csv(METRIC_DIR / "table6_flow_efficiency.csv")


def _build_table8() -> pd.DataFrame:
    return read_csv(METRIC_DIR / "table8_grayzone.csv")


def _build_table10() -> pd.DataFrame:
    df = read_csv(METRIC_DIR / "table10_operational_strategy.csv").copy()
    df["Call Volume Summary"] = df.apply(
        lambda r: f"{r['Call Volume Summary']}, time {r['Inference Time (s)']} s",
        axis=1,
    )
    return df[["Strategy", "Invocation Structure", "Call Volume Summary", "Operational Significance"]]


def _build_appendix_q_table() -> pd.DataFrame:
    for path in [
        APPENDIX_DIR / "table_a2_q_sweep_base_models.csv",
        APPENDIX_DIR / "table_d1_q_sweep_base_models.csv",
    ]:
        if path.exists():
            return read_csv(path)
    raise FileNotFoundError("No appendix q-sweep base-model table found.")


def _build_appendix_a0_table() -> pd.DataFrame:
    return read_csv(APPENDIX_DIR / "table_a0_baseline_descriptions.csv")


def _build_appendix_a1_table() -> pd.DataFrame:
    return read_csv(APPENDIX_DIR / "table_a1_supplementary_experimental_results.csv")


def _build_appendix_e_distribution_table() -> pd.DataFrame:
    df = read_csv(APPENDIX_DIR / "table_e1_distribution_shift_summary.csv").copy()
    renamed = pd.DataFrame(
        {
            "Comparison": df["comparison"],
            "Avg |Δmean|": df["avg_mean_abs_diff"].map(lambda x: f"{float(x):.4f}"),
            "Avg KS": df["avg_ks_stat"].map(lambda x: f"{float(x):.4f}"),
            "KS Reject @0.05": df["ks_reject_ratio_0.05"].map(lambda x: f"{float(x):.4f}"),
            "MMD (RBF)": df["mmd_rbf"].map(lambda x: f"{float(x):.4f}"),
        }
    )
    return renamed


def _build_appendix_g_seed_detail_table() -> pd.DataFrame:
    df = read_csv(APPENDIX_DIR / "table_g1_seed_variation_detail.csv").copy()
    for col in ["Ave. F1", "PRR", "Worst-Case Recall (P5)", "Instability (Var)"]:
        df[col] = df[col].map(lambda x: f"{float(x):.4f}")
    df["Seed Number"] = df["Seed Number"].astype(int).astype(str)
    return df[["Seed Number", "Method", "Ave. F1", "PRR", "Worst-Case Recall (P5)", "Instability (Var)"]]


def _maybe_write_table(table, builder) -> None:
    try:
        df = builder()
    except FileNotFoundError:
        return
    if df is None or df.empty:
        return
    _write_dataframe_to_table(table, df)


def _build_appendix_strategy_table() -> pd.DataFrame:
    q_table = read_csv(METRIC_DIR / "table3_q_sweep.csv").copy()
    selected_q = read_selected_q(DEFAULT_Q)
    interpretations = []
    q_values = q_table["q"].tolist()
    for q in q_values:
        if np.isclose(q, q_values[0]):
            label = "Economy Mode"
        elif np.isclose(q, selected_q):
            label = "Selected Mode"
        elif np.isclose(q, q_values[-1]):
            label = "Safety-First Mode"
        else:
            label = "Intermediate Mode"
        interpretations.append(label)
    q_table["Operational Interpretation"] = interpretations
    q_table["Cost"] = q_table["Cost (USD)"]
    q_table["LLM Call Rate(Total)"] = q_table["Call Rate"]
    return q_table[["q", "Gray Ratio", "LLM Call Rate(Total)", "Cost", "Operational Interpretation"]]


def _replace_images(docx_path: Path) -> None:
    replacements = {
        "word/media/image1.png": FIG_DIR / "figure_methodology_evidence.png",
        "word/media/image3.png": FIG_DIR / "figure2_qsweep_elbow.png",
        "word/media/image4.png": FIG_DIR / "figure3_callrate_vs_f1.png",
        "word/media/image5.png": FIG_DIR / "figure4_performance_drop.png",
        "word/media/image6.png": FIG_DIR / "figure5_prediction_flips.png",
        "word/media/image8.png": FIG_DIR / "figure6_cost_stability_pareto.png",
        "word/media/image9.png": APPENDIX_DIR / "figure_f1_graphad_visual_proof.png",
    }
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        with ZipFile(docx_path) as zf:
            zf.extractall(tmpdir_path)
        for internal, src in replacements.items():
            if src.exists():
                shutil.copy2(src, tmpdir_path / internal)
        rebuilt = docx_path.with_suffix(".tmp.docx")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as zf:
            for path in sorted(tmpdir_path.rglob("*")):
                if path.is_file():
                    zf.write(path, path.relative_to(tmpdir_path))
        rebuilt.replace(docx_path)


def _resize_inline_shape(doc, *, partname: str, width_inches: float, height_inches: float) -> None:
    for shape in doc.inline_shapes:
        inline = shape._inline
        blip = inline.graphic.graphicData.pic.blipFill.blip
        rid = blip.embed
        part = doc.part.related_parts[rid]
        if str(part.partname) == partname:
            shape.width = Inches(width_inches)
            shape.height = Inches(height_inches)
            return


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    selected_q = read_selected_q(DEFAULT_Q)
    seed_count = get_seed_count()
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)

    doc = Document(TARGET_DOCX)
    table2_df = _build_table2(selected_q)
    table3_df = _build_table3()
    table5_df = _build_table5(selected_q)
    table2_full = read_csv(METRIC_DIR / "table2_robustness.csv")

    _write_dataframe_to_table(doc.tables[0], _build_table1())
    _write_dataframe_to_table(doc.tables[1], table2_df)
    _write_dataframe_to_table(doc.tables[2], table3_df)
    _write_dataframe_to_table(doc.tables[4], table5_df)
    _write_dataframe_to_table(doc.tables[5], _build_table6())
    _write_dataframe_to_table(doc.tables[6], _build_table8())
    _write_dataframe_to_table(doc.tables[8], _build_table10())
    _maybe_write_table(doc.tables[9], _build_appendix_a0_table)
    _maybe_write_table(doc.tables[10], _build_appendix_a1_table)
    _maybe_write_table(doc.tables[12], _build_appendix_q_table)
    _maybe_write_table(doc.tables[13], _build_appendix_e_distribution_table)
    _maybe_write_table(doc.tables[14], _build_appendix_g_seed_detail_table)
    _write_dataframe_to_table(doc.tables[15], _build_appendix_strategy_table())

    flow = read_csv(METRIC_DIR / "table6_flow_efficiency.csv")
    confident = flow.iloc[1]
    shortcut = flow.iloc[2]
    llm = flow.iloc[3]
    grayzone = read_csv(METRIC_DIR / "table8_grayzone.csv")
    q_sweep = read_csv(METRIC_DIR / "table3_q_sweep.csv").sort_values("q").reset_index(drop=True)
    q_first = float(q_sweep["q"].min())
    q_last = float(q_sweep["q"].max())
    selected_q_row = q_sweep[np.isclose(q_sweep["q"], selected_q)].iloc[0]
    next_q_rows = q_sweep[q_sweep["q"] > selected_q].sort_values("q").reset_index(drop=True)
    next_q_row = next_q_rows.iloc[0] if not next_q_rows.empty else None
    selected_gray = grayzone[np.isclose(grayzone["q"], selected_q)]
    utar_gray = selected_gray[selected_gray["Method"] == "UTAR (No-LLM)"].iloc[0]
    routing_gray = selected_gray[selected_gray["Method"] == "Selective Routing"].iloc[0]
    utar_gray_all = grayzone[grayzone["Method"] == "UTAR (No-LLM)"].sort_values("q").reset_index(drop=True)
    gray_first = utar_gray_all.iloc[0]["Gray Ratio"]
    gray_last = utar_gray_all.iloc[-1]["Gray Ratio"]
    utar_table2 = table2_full[table2_full["Method"] == "UTAR (Proposed)"].iloc[0]
    avg_ensemble_table2 = table2_full[table2_full["Method"] == "Avg. Ensemble (Base)"].iloc[0]
    moderntcn_table2 = table2_full[table2_full["Method"] == "ModernTCN (Base)"].iloc[0]
    full_framework = table5_df[table5_df["Configuration"] == "Full Framework (UTAR)"].iloc[0]
    wo_selective = table5_df[table5_df["Configuration"] == "w/o Selective LLM"].iloc[0]
    wo_routing = table5_df[table5_df["Configuration"] == "w/o Gray-Zone (Routing)"].iloc[0]
    wo_graph = table5_df[table5_df["Configuration"] == "w/o Graph Smoothing"].iloc[0]
    wo_entropy = table5_df[table5_df["Configuration"] == "w/o Ensemble Entropy"].iloc[0]
    cost_selective = _selected_summary_row("cost", "selective", selected_q)
    cost_no_llm = _selected_summary_row("cost", "no_llm", selected_q)
    cost_full_llm = _selected_summary_row("cost", "full_llm", selected_q)
    next_q_sentence = ""
    if next_q_row is not None:
        next_q_sentence = (
            f" 더 높은 q 후보들은 F1 개선 없이 Call Rate와 Cost만 증가하는 경향을 보여, "
            f"선택된 q={selected_q:.2f}의 비용-성능 균형을 넘어서지 못했다."
        )

    replacements = {
        3: "Heterogeneous Ensemble Entropy Filter 기반의 선택적 LLM 라우팅을 통해 full-LLM upper bound 대비 운영 비용을 91.8% 절감.",
        8: "첫째, 구조적 성능 회복력(Resilience)이다. Quantile 기반의 Gray-Zone 메커니즘을 통해 시프트 발생 시 나타나는 결정 경계의 모호성을 식별하고, 불확실성이 높은 샘플만을 선별적으로 상위 추론 모듈(Selective LLM)로 전달함으로써 TDS 환경에서도 Ave. F1 0.8835와 PRR 0.9807을 달성하였다. 둘째, 실무적 비용 효율성(Cost-Efficiency)이다. 모든 데이터를 고비용 연산 모델로 처리하는 대신 계층적 라우팅 경로를 설계하여, cost split 기준 보조 모듈 호출 비용을 $0.0088로 유지하면서 full-LLM upper bound의 $0.1073 대비 91.8% 절감하였다. 셋째, 구조적 진단 가시성(Decision Support)이다. GraphAD+ 모듈을 통해 공정의 물리적 연결성을 고려한 그래프 기반 스무딩 점수를 산출하고, 이를 시각화된 대시보드와 후보군 정보로 제공함으로써 운영자가 시프트 환경에서도 일관되고 설명 가능한 진단 근거를 확보하도록 지원한다.",
        18: "• 제조 공정 내 시간적 분포 변화가 base detector의 PRR과 경계 안정성을 약화시킴을 실증\n• Quantile 기반 Gray-Zone 라우팅 구조를 통해 TDS에서 Ave. F1 0.8835와 PRR 0.9807 달성\n• Avg. Ensemble 대비 Worst-Case Recall (P5)을 0.0582에서 0.1758로 향상\n• 선택적 LLM 전략으로 full-LLM upper bound 대비 운영 비용 91.8% 절감\n• 구조적 GraphAD+ evidence와 불확실성 라우팅을 결합한 실무형 diagnosis pipeline 제시",
        35: "불확실성 격리 기반의 고강건성 UTAR 라우팅 설계 단순 앙상블 기법의 한계를 극복하기 위해, TDS 환경에서 급증하는 결정 경계의 모호성(Gray-Zone)을 실시간으로 식별하고 격리하는 UTAR 라우팅 아키텍처를 설계하였다. 이를 통해 Avg. Ensemble 대비 PRR을 0.9310에서 0.9807로 높였고, Worst-Case Recall (P5)을 0.0582에서 0.1758로 향상시켜 shift 구간의 safety floor를 강화하였다.",
        37: "성능 손실 없는 비용 최적화(Cost-Efficient High Performance) 전략 LLM의 강력한 추론 능력을 전역적으로 사용하는 대신, 고난도 불확실 구간에만 집중 투여하는 선택적 호출 전략을 제안하였다. 그 결과, selective routing은 cost split에서 $0.0088의 비용과 42.7108 s의 latency로 동작하며, full-LLM upper bound의 $0.1073 대비 운영 비용을 91.8% 절감하였다.",
        52: "위 연구들을 종합하면, 기존 접근법은 (i) 정적인 IID 가정 기반의 탐지 최적화, (ii) 고비용 Full-Engine LLM 구조, (iii) 물리적 공정 구조가 배제된 통계적 원인 분석이라는 공통적 한계를 지닌다. 특히 시간적 분포 변화 상황에서 의사결정 안정성과 LLM 운영 비용 통제를 동시에 만족하는 통합 프레임워크는 충분히 제시되지 않았다. 이에 본 연구는 UTAR 기반의 불확실성 라우팅, GraphAD+를 활용한 구조적 후보 생성, 그리고 선택적 LLM 호출을 결합한 운영형 진단 프레임워크를 제안한다.\n\n3. Proposed Methodology: Robust Anomaly Detection and Diagnosis Framework\n본 절에서는 시간적 분포 변화(Temporal Distribution Shift, TDS) 환경에서도 의사결정 안정성을 유지하기 위해 설계한 UTAR-GraphAD+ 프레임워크를 설명한다. Figure 1과 같이 본 프레임워크는 탐지(Detection)와 진단(Diagnosis)을 계층적으로 분리하고, 불확실성이 높은 샘플만 상위 추론 경로로 라우팅하여 성능과 비용을 함께 관리한다.",
        53: "Figure 1은 제안 프레임워크의 필요성을 유발하는 세 가지 데이터 현상을 시각적으로 요약한다.",
        54: "첫째, 분포 중첩(gray-zone emergence)이다. IID validation 환경에서는 정상과 이상 score 분포가 비교적 뚜렷하게 분리되지만, TDS 환경에서는 두 분포가 decision threshold 주변으로 재집중되며 겹치는 구간이 커진다.",
        55: "이 겹치는 구간은 단일 임계값 기반의 이진 판정이 가장 취약해지는 영역이며, 실제로 UTAR는 이 경계 구간을 gray-zone으로 분리하여 별도의 라우팅 대상으로 취급한다.",
        56: "즉 Figure 1(a)의 핵심 메시지는 '분포가 겹치는 운영 구간에서는 고정 threshold만으로 안정적인 판단을 기대하기 어렵다'는 점이다.",
        57: "둘째, decision-boundary margin concentration이다.",
        58: "TDS 환경에서는 anomaly score가 임계치 주변에 밀집하여, 아주 작은 노이즈나 확률 변화만으로도 판정이 정상과 이상 사이를 오갈 수 있다.",
        59: "이는 평균 성능이 유지되더라도 운영 시점의 decision volatility가 커질 수 있음을 의미하며, 본 연구가 margin-q를 통해 gray-zone 폭을 직접 제어하는 이유이기도 하다.",
        60: "셋째, inter-model discrepancy이다.",
        61: "정적 모델(RF/XGB)과 동적 모델(ModernTCN)이 같은 샘플을 평가하더라도, IID에서는 score가 상대적으로 일관되게 정렬되지만 TDS에서는 대각선에서 벗어난 불일치 샘플이 증가한다.",
        62: "이러한 불일치는 단순 앙상블 평균으로는 충분히 해소되지 않으며, 오히려 어떤 샘플을 추가 검토 대상으로 삼아야 하는지에 대한 신호가 된다.",
        63: "UTAR는 바로 이 불일치와 경계 부근 불확실성을 함께 사용하여, 운영 리스크가 큰 샘플만 선택적으로 LLM 경로로 상승시킨다.",
        64: "따라서 Figure 1은 제안 방법의 목적이 평균 정확도 개선만이 아니라, TDS에서 증가하는 경계 샘플의 불안정성을 구조적으로 통제하는 데 있음을 보여준다.",
        66: "Figure 1. Proposed Robust Diagnosis Framework: Integrating Uncertainty-aware Routing and Structure-aware GraphAD+ with Selective LLM. The figure summarizes three operational symptoms under temporal shift - distribution overlap, boundary-margin concentration, and inter-model discrepancy - and motivates the staged UTAR-GraphAD+ pipeline used throughout this study.\n\n3.1 Design Objectives: Stability under Temporal Distribution Shift (TDS)",
        107: "4장 실험에서는 이 안정성 명제를 단일 정확도 비교가 아니라 PRR, Worst-Case Recall (P5), Instability, 그리고 Prediction Flip 분석으로 검증한다. 즉 핵심 주장은 'UTAR가 평균 점수만 높다'가 아니라, shift 구간에서도 의사결정 하한선과 운영 안정성을 더 잘 유지한다는 점이다.",
        110: "Multi-faceted Robust Scoring",
        111: "공격 및 이상 구간에서 발생하는 센서 데이터의 변동성을 다각도로 포착하기 위해, 본 연구에서는 정상 구간의 median과 robust scale을 기준으로 세 가지 sensor-level score를 산출한다. 이는 평균과 표준편차 기반 점수가 이상치에 의해 과도하게 왜곡되는 문제를 완화하기 위함이다 (Leys et al., 2013).",
        112: "1. Robust Z-Score (): 현재 센서값이 정상 기준에서 얼마나 이탈했는지를 측정한다 (Eq. 6).",
        114: "2. Trend Score (): 인접 시점 간 변화량을 이용해 onset 이후의 방향성 변화와 slope 변화를 포착한다 (Eq. 7).",
        116: "3. Fluctuation Score (): curvature 기반 변동성을 이용해 비정상적인 진동과 turbulence를 포착한다 (Eq. 8).",
        118: "위 세 가지 robust score는 가중합을 통해 통합 node score를 구성한다(Eq. 9). 본 구현에서는 λ_z = 0.4, λ_tr = 0.3, λ_fl = 0.3을 사용하여 instantaneous deviation 신호에 가장 큰 비중을 두었다.",
        119: "여기서 각 score는 각각 순간 편차, 방향성 변화, 비정상 진동을 반영하며, 세 계수의 합은 1이 되도록 정규화된다. 이 설정은 Appendix H의 하이퍼파라미터 근거표와 configs/train_graphad.yaml의 current configuration을 따른다.",
        120: "즉 GraphAD+의 통합 점수는 단일 센서의 급격한 이탈을 우선적으로 반영하되, 추세와 변동성 정보를 함께 고려하여 구조적으로 일관된 이상 징후를 상위 후보로 올리도록 설계되었다.",
        123: "여기서 N(v)는 공정 그래프 상에서 노드 v와 연결된 인접 센서 집합이며, α는 개별 센서의 raw anomaly evidence와 이웃 기반 neighborhood support를 얼마나 혼합할지를 결정하는 스무딩 계수이다. 본 구현에서는 α = 0.3을 사용하여 단일 spike 노이즈를 완화하면서도 특정 센서의 고유 징후가 과도하게 희석되지 않도록 균형을 맞췄다.",
        124: "3.4.3 Evaluation and Hybrid Re-ranking (Auxiliary)\n실험적 비교를 위해 GraphAD+ 점수와 LLM ranking score를 결합한 auxiliary hybrid reranking 식(Eq. 11)도 함께 검토하였다. 다만 본 논문의 TE 탐지 표(Table 2, 3, 5)에 보고한 핵심 UTAR 결과는 이 보조 식이 아니라, Eq. 12의 q 기반 selective routing 정책에서 얻어진 운영 성능이다. 마지막으로 본 프레임워크의 비용 효율성은 호출 비율(Call Rate), LLM cost, 그리고 latency를 함께 사용하여 정량화한다.",
        203: "4.5.3 Diagnosis Stability via GraphAD+\nGraphAD+의 역할은 detection AUC를 직접 끌어올리는 것보다, 구조적 후보군을 안정화하여 후속 진단 단계가 일관된 증거를 보도록 만드는 데 있다. 따라서 본 연구는 GraphAD+를 root-cause candidate ranking과 prompt evidence quality를 높이는 모듈로 해석하며, detection boost만으로 평가하지 않는다.",
        231: "Figure 4는 IID와 TDS 사이의 성능 저하량(performance drop)을 모델별로 비교한 시각화이다.",
        232: "이 그림의 목적은 단일 최고 점수를 강조하는 것이 아니라, temporal shift 이후 Ave. F1, PRR, Worst-Case Recall (P5)과 같은 운영 지표가 얼마나 안정적으로 유지되는지를 모델별로 비교하는 데 있다.",
        233: "따라서 Figure 4는 UTAR가 base detectors 대비 성능 붕괴를 얼마나 완만하게 통제하는지, 즉 operating point 유지력 측면에서 어떤 이점을 가지는지를 직관적으로 보여주는 보조 자료로 해석해야 한다.",
        411: "Appendix H는 본 구현에서 사용한 핵심 하이퍼파라미터와 그 선택 근거를 정리한 섹션이다. 아래 설명은 현재 코드(configs/*.yaml)와 Appendix H의 근거표를 기준으로 작성하였다.",
        414: "설정값: λ_z = 0.4, λ_tr = 0.3, λ_fl = 0.3.",
        415: "근거: λ_z는 instantaneous deviation을 가장 직접적으로 반영하므로 가장 큰 가중치를 부여하였다.",
        416: "반면 λ_tr와 λ_fl는 onset 이후의 방향성 변화와 비정상 진동을 보강하는 항으로 두어, 단일 시점 이상치에 과도하게 끌리지 않도록 균형을 맞췄다.",
        417: "이 선택은 Appendix H의 candidate space와 current configuration을 따르며, 구조적 후보 생성 단계에서 local spike sensitivity와 ranking stability를 함께 확보하기 위한 설정이다.",
        419: "그래프 기반 스무딩 단계는 상관관계 기반 공정 그래프에서 인접 센서 정보를 얼마나 반영할지를 결정하는 계수이다.",
        420: "설정값: alpha = 0.3.",
        421: "근거: alpha가 0이면 raw score에만 의존하여 단일 센서 spike에 민감해지고, alpha가 지나치게 크면 개별 센서의 고유 징후가 이웃 평균에 묻혀 해상도가 낮아진다.",
        422: "따라서 본 연구는 local evidence와 neighborhood support를 절충하는 중간값 0.3을 사용하여, GraphAD+의 구조적 일관성과 개별 센서 해상도를 함께 유지하였다.",
        423: "이 값 역시 configs/train_graphad.yaml과 Appendix H의 candidate space([0.2, 0.3])에 기반한 current configuration이다.",
        424: "즉 Process Graph Smoothing은 raw sensor data 전처리 제거/추가의 문제가 아니라, GraphAD가 생성한 sensor-level anomaly score를 공정 토폴로지 위에서 안정화하는 단계로 이해해야 한다.",
        425: "3. Shortcut Quantiles for Selective Routing",
        426: "설정값: entropy_shortcut_quantile = 0.8, discrepancy_shortcut_quantile = 0.8.",
        427: "근거: shortcut 기준이 너무 낮으면 애매한 샘플이 조기에 확정되어 recall 손실이 커지고, 너무 높으면 대부분의 샘플이 LLM 경로로 넘어가 비용과 지연이 급증한다.",
        428: "본 연구는 두 기준을 동일한 0.8로 두어, 고신뢰 샘플은 base routing으로 먼저 결정하고 불확실 샘플만 selective escalation 하도록 설계하였다.",
        429: "이 설정은 Appendix H의 candidate space([0.8, 0.9])에서 선택된 current configuration이며, cost split 비교에서도 호출량 통제와 경계 샘플 보호의 균형을 뒷받침한다.",
        430: "즉 shortcut quantile은 최고 정확도만이 아니라 cost-performance-stability의 운영 균형점을 만들기 위한 파라미터이다.",
        431: "4. Uncertainty Threshold (Margin-q)",
        432: f"설정값: q = {selected_q:.2f}.",
        433: f"근거: Pareto efficiency 기준으로 q={selected_q:.2f} 지점이 F1, Recall, Worst-Case Recall (P5), Call Rate, Cost의 균형이 가장 우수한 선택 지점으로 확인되었다.",
        434: "q는 gray-zone 폭을 직접 조절하는 하이퍼파라미터로, 값이 커질수록 더 많은 경계 샘플이 상위 라우팅 경로로 escalation된다. 본 실험에서는 q=0.40이 추가 비용 증가 없이 가장 높은 F1을 제공하여 최종 operating point로 채택되었다.",
        435: "이 선택은 본문 q-sweep 결과(Figure 3, Table 3, Table 8)와 Appendix A2/Table 16의 supplementary evidence와 일관된다.",
        436: "5. Sigmoid Gain",
        437: "설정값: sigmoid gain k = 5.0.",
        438: "근거: k는 decision boundary 부근에서 temporal weight가 얼마나 급하게 변하는지를 결정한다. 너무 낮으면 gray-zone과 confident zone의 분리가 완만해지고, 너무 높으면 작은 확률 변동에도 routing weight가 과민하게 반응한다.",
        439: "본 연구는 current configuration인 k = 5.0을 사용하여 경계 부근 샘플에만 선택적 가중치를 집중하고, 확실한 샘플은 base detector 경로에 남기도록 설계하였다. 또한 GraphAD+는 top_k = 5를 사용하여 prompt에 전달되는 구조적 후보 센서 수를 제한함으로써, 핵심 후보 집중과 prompt 길이 통제를 동시에 달성한다.",
        440: "Table H1. Summary of LLM Call Rate and Operational Strategy.",
        441: "Table H2. Robustness Analysis across Random Seeds.",
        442: "Table H2는 Table 2의 평균 뒤에 숨은 seed-level 변동을 보여주며, UTAR가 전반적으로 가장 높은 F1과 안정적인 PRR을 유지한다는 점을 확인하게 한다.",
        444: "Table I1. NDSS: Root-Cause Candidate Performance on 286 Scenarios (LLM Calls = 286).",
        445: "이 표는 NDSS 진단 시나리오에서 후보 생성 및 reranking 전략의 상대적 장단점을 비교하기 위한 보조 분석으로 해석해야 한다.",
        292: "안전 하한선 확보: UTAR는 TDS 환경에서 Recall 0.8207과 Worst-Case Recall (P5) 0.1758을 기록하여, Avg. Ensemble의 0.7791과 0.0582보다 더 강한 safety floor를 제공한다.",
        293: "비용 효율 극대화: cost split에서 selective routing은 $0.0088의 비용과 call rate 0.0825로 동작하여, full-LLM upper bound의 $0.1073 대비 91.8% 낮은 비용으로 운영 가능성을 보여준다.",
        317: "LLM을 전역적 보정 도구(Global Optimizer)가 아닌 선택적 보조 모듈(Uncertainty-aware Auxiliary Module)로 사용해야 한다는 점을 실증했다. 전면적 호출은 비용 대비 이득이 제한적일 뿐 아니라 변동성을 확대할 위험이 있다. 반면 본 연구의 Selective Routing은 cost split 기준 $0.0088의 비용으로 full-LLM upper bound 대비 91.8% 낮은 비용을 유지하면서도 경계 샘플 보강 판단을 수행한다.",
        145: "이상 탐지 성능 평가는 공개 Tennessee Eastman Process (TEP) 시뮬레이션 변형본(Kaggle 버전)을 사용하였다. TEP는 복잡한 변수 간 상관관계와 비선형성을 포함하므로 산업 공정 모니터링 연구의 대표 벤치마크로 널리 사용된다 (Downs & Vogel, 1993). 본 구현은 4개의 원천 CSV(TEP_FaultFree_Training.csv, TEP_Faulty_Training.csv, TEP_FaultFree_Testing.csv, TEP_Faulty_Testing.csv)를 병합하고, 각 run을 500 step 시계열로 정렬한 뒤 41개 측정 변수(XMEAS)와 11개 조작 변수(XMV)를 사용한다. Fault 0은 정상, Fault 1-20은 고장으로 두고, 최종 탐지 문제는 정상-이상 이진 분류로 재구성하였다.",
        146: "중요하게도 본 연구의 데이터 분할은 임의 row split이 아니라 run-preserving temporal split으로 구현되었다. 즉 source_file, fault_id, run_id를 보존한 채 학습, 검증, 평가 run을 먼저 샘플링하고, 이후에 시간 축 기반 phase label을 부여하였다. 이 절차는 build_split.py에서 직접 구현되어 있으며, 동일 run 내부의 시간적 연속성을 깨지 않도록 구성하였다.",
        147: "Temporal Distribution Shift는 고장 발생 시점(onset)과 전이 구간(transition length)을 명시적으로 정의하여 재현하였다. 코드에서는 onset metadata를 source_file과 fault_id 단위로 먼저 생성하고, 이를 각 row에 merge한 뒤 sample_idx 기준으로 phase를 판정한다. 정상 run은 onset이 없는 것으로 처리하기 위해 onset_step을 매우 큰 값으로 두고, 고장 학습 run은 onset_step = 20, 고장 테스트 run은 onset_step = 160, transition_len = 50으로 설정하였다.",
        148: "Phase는 다음과 같이 구현된다. normal은 fault_id = 0인 fault-free run 전체 구간이다. pre는 fault_id != 0인 고장 run이지만 sample_idx < onset_step인 구간으로, 고장 run이더라도 아직 분포 변화가 시작되지 않은 구간이다. transition은 fault_id != 0이면서 onset_step <= sample_idx < onset_step + transition_len을 만족하는 구간으로, 고장 주입 직후 분포가 이동하기 시작하는 초기 전이 영역이다. post_shift는 fault_id != 0이면서 sample_idx >= onset_step + transition_len인 구간으로, fault 이후 새로운 운영 분포가 형성된 뒤의 안정화 영역이다.",
        149: "라벨 부여도 phase 정의와 직접 연결된다. 구현상 y는 먼저 0으로 초기화한 뒤, phase가 transition 또는 post_shift인 경우에만 y = 1로 바꾼다. 따라서 normal과 pre는 모두 negative class(0), transition과 post_shift는 positive class(1)로 학습 및 평가된다. 이 설계는 고장 run의 onset 이전 구간을 정상에 가까운 운영 상태로 취급하고, 실제로 탐지해야 하는 분포 변화는 onset 이후부터라는 문제 정의를 코드 수준에서 반영한 것이다.",
        150: "학습과 검증은 training-domain run에서만 구성하였다. 현재 설정(configs/split.yaml) 기준으로 train은 정상 95개 run과 각 fault당 6개 run을 사용하여 총 215개 run unit, 107,500개 row를 만든 뒤, RF와 XGBoost 학습용으로는 class imbalance를 줄이기 위해 99,560개 row의 1:1 balanced train set을 별도로 생성하였다. Validation은 정상 8개 run과 각 fault당 1개 run으로 구성되며, 총 28개 run unit과 14,000개 row를 포함한다.",
        151: "평가용 테스트는 testing-domain run에서만 구성하였다. 먼저 정상 20개 run과 각 fault당 5개 run을 선택하여 총 120개 run unit의 contiguous test set(test_full_tcn, 115,200 rows)을 만들고, 여기에는 pre 구간도 유지하였다. 이후 공통 평가용 row-level split에서는 pre를 제외한 normal, transition, post_shift만 남긴 뒤, robustness와 q-sweep selection을 위한 test_main 4,000 rows와 운영 비용 비교를 위한 test_cost 400 rows를 추가로 샘플링하였다. test_main은 normal 30%, transition 20%, post_shift 50%(= 1,200 / 800 / 2,000 rows), test_cost는 normal 30%, transition 35%, post_shift 35%(= 120 / 140 / 140 rows) 비율로 구성하였다.",
        152: "이처럼 pre 구간을 test_main과 test_cost에서는 제외하고 test_full_tcn에는 유지한 이유는, 보고용 평가는 실제 운영에서 중요한 transition과 post_shift 강건성에 집중하되, ModernTCN과 GraphAD+는 full-run 문맥을 이용해 추론하도록 하기 위함이다. 실제로 window 기반 모델 입력은 build_windows.py에서 window_size = 50, stride = 1로 생성되며, source_file, fault_id, run_id가 동일하고 sample_idx가 연속인 경우에만 window를 만든다. 따라서 ModernTCN과 GraphAD+는 full-run contiguous context를 사용하고, 최종 표와 q 선택은 compact row-level split(test_main과 test_cost)에서 수행된다.",
        158: f"Detection 모델 설정은 표 1과 같다. 모든 실험은 결과의 신뢰성을 확보하기 위해 {seed_count}개의 서로 다른 랜덤 시드(Random Seeds)를 사용하여 평균과 표준편차를 도출하였다 (Kuncheva & Whitaker, 2003). 베이스라인 모델인 ModernTCN, RF, XGBoost는 선행 연구에서 검증된 하이퍼파라미터 세팅을 준수하였다(Bai et al., 2018; Chen & Guestrin, 2016). Table 1. Model Configuration Details",
        171: "Performance Retention Rate (PRR): validation 환경에서 얻은 recall 대비 temporal shift가 발생한 테스트 환경의 recall 비율로 정의한다. 본 구현에서는 PRR = Recall_test / Recall_val을 사용하며, 값이 1에 가까울수록 시프트 이후에도 재현율이 잘 유지됨을 의미한다 (Eq. 15).",
        172: "Worst-Case Recall (P5): 각 run 내부에서 길이 50의 sliding window를 이동시키며 anomaly가 포함된 window의 recall을 계산한 뒤, 그 분포의 하위 5% 분위수를 취한 값이다. 즉 평균 recall이 아니라 실제 운영 중 가장 어려운 구간들에서 시스템이 보장하는 최소 안전 하한선을 나타낸다 (Eq. 16).",
        195: "Table 2의 모든 수치는 10개 랜덤 시드에 대한 평균 ± 표준편차로 보고하였다. 본문 해석은 통계적 유의성 표기보다 effect size와 운영적 의미에 초점을 맞추었다.",
        196: f"4.5.2 Economic Feasibility and Selective Logic\nq={selected_q:.2f}로 선택된 selective routing은 cost split에서 call rate {cost_selective['llm_call_rate_mean']:.4f}, cost ${cost_selective['cost_usd_mean']:.4f}, inference time {cost_selective['total_latency_ms_mean']/1000.0:.4f} s를 기록하였다. 이는 full-LLM upper bound의 cost ${cost_full_llm['cost_usd_mean']:.4f}와 비교해 약 {(1.0 - cost_selective['cost_usd_mean'] / cost_full_llm['cost_usd_mean']) * 100.0:.1f}% 낮은 수준이며, no-LLM baseline의 비용-지연 이점과 full-LLM의 정보 활용 사이에서 실용적인 균형점을 제공한다.",
        210: "본 실험은 one-at-a-time ablation 설계로 수행하였다. 즉 Full Framework를 기준으로 각 구성요소를 하나씩 제거하거나 대체하여, PRR, Worst-Case Recall (P5), Instability, 그리고 비용-지연 변화가 어떤 방식으로 달라지는지를 동일한 test_main 분할에서 비교하였다.",
        212: f"분석: Table 5 기준으로 선택적 LLM을 제거하면 Ave. F1이 {wo_selective['Ave. F1']}, PRR이 {wo_selective['PRR']}, Worst-Case Recall (P5)이 {wo_selective['Worst-Case Recall (P5)']}로 낮아져, Full Framework의 {full_framework['Ave. F1']}, {full_framework['PRR']}, {full_framework['Worst-Case Recall (P5)']} 대비 경계 샘플 방어력이 약해진다. Gray-Zone 라우팅 자체를 제거한 경우에도 F1과 PRR이 {wo_routing['Ave. F1']}, {wo_routing['PRR']}로 더 낮아진다. 반면 graph smoothing 제거는 현재 설정에서 Full Framework와 거의 동일한 수치를 보이므로, 본 실험에서는 graph smoothing의 기여가 크기보다는 제한적이라고 해석하는 것이 타당하다. 가장 큰 운영 차이는 엔트로피 필터를 제거했을 때 나타나며, 이 경우 지연이 {wo_entropy['Inference Time (s)']}까지 급증한다.",
        213: "Ablation 설정은 아래와 같이 정리된다.",
        218: "관찰 포인트: LLM을 제거했을 때 Gray-Zone 경계 샘플에서 Worst-Case Recall이 얼마나 감소하는지 확인한다.",
        221: "관찰 포인트: 라우팅 자체를 제거했을 때 F1과 PRR이 Full Framework 대비 얼마나 감소하는지 확인한다.",
        222: "3. ④ w/o Graph Smoothing (GraphAD 스무딩 제거)",
        223: "방법: GraphAD의 smoothed structural context 대신 raw GraphAD context를 입력으로 사용합니다.",
        224: "관찰 포인트: graph smoothing이 경계 샘플의 판단 일관성과 안정성에 주는 추가 이득을 확인한다.",
        234: f"실험 결과, 단순 앙상블(Simple Ensemble) 모델은 TDS 환경에서 Ave. F1 {avg_ensemble_table2['Ave. F1']}과 PRR {avg_ensemble_table2['PRR']}에 머물렀다. 반면 UTAR는 Ave. F1 {utar_table2['Ave. F1']}과 PRR {utar_table2['PRR']}을 기록하여 시프트 이후 성능 유지력이 더 높았다. 특히 UTAR는 Worst-Case Recall (P5) {utar_table2['Worst-Case Recall (P5)']}를 보여 평균 확률 앙상블의 {avg_ensemble_table2['Worst-Case Recall (P5)']}보다 더 높은 안전 하한선을 확보하였다. 이는 UTAR가 단순 평균 앙상블보다 경계 샘플을 더 안정적으로 처리함을 의미한다.",
        281: f'Ensemble Entropy Filter의 역할 (Stage 1): "Table 7에서 알 수 있듯이, 전체 라우팅 테스트 샘플 4,000개 중 {confident["Filtered / Decided"].split()[0]}개가 Confident Zone에서 즉시 처리되었다. 이는 명확한 상태를 가벼운 모델이 우선적으로 판정함으로써 시스템 전체의 부하를 크게 줄였음을 의미한다."',
        282: f'Gray-Zone Routing의 역할 (Stage 2): "남은 {shortcut["Input Samples"].split()[0]}개의 Gray-Zone 샘플 중 엔트로피 shortcut으로 {shortcut["Filtered / Decided"].split()[0]}개를 먼저 확정하고, 오직 최종 불확실 샘플 {llm["Filtered / Decided"].split()[0]}개만 LLM으로 라우팅하였다."',
        287: f"1. Gray-Zone의 적응적 확장성과 모호성 제어 실험 결과, 의사결정 마진 q가 {q_first:.2f}에서 {q_last:.2f}로 증가함에 따라 Gray-Zone 비율은 {gray_first}에서 {gray_last}로 상승하였다. 이는 시스템이 판단 유보 영역을 점진적으로 넓히며 모호한 샘플을 상위 추론 영역으로 더 많이 포섭하고 있음을 보여준다.",
        288: f'2. 하이브리드 라우팅의 질적 우월성 입증 모든 q 구간에서 제안된 Routing 전략은 Gray-Zone 내부 성능에서 단독 UTAR보다 높은 F1을 보였다. 특히 q={selected_q:.2f}에서 Gray-Zone F1은 UTAR {utar_gray["F1 (Gray)"]} 대비 Routing {routing_gray["F1 (Gray)"]}로 개선되었고, Gray-Zone recall도 {utar_gray["Rec (Gray)"]}에서 {routing_gray["Rec (Gray)"]}로 상승하였다. 반면 AUC는 {utar_gray["AUC (Gray)"]}에서 {routing_gray["AUC (Gray)"]}로 소폭 낮아져, 본 결과는 전역 ranking 성능 향상보다는 경계 샘플에 대한 recall-oriented correction 효과로 해석하는 것이 타당하다.',
        289: f"3. 한계 효용 및 최적 운영 지점(Sweet Spot) 도출 q-sweep 결과에서 선택된 운영 지점은 q={selected_q:.2f}이며, 이 지점에서 F1과 호출 비용의 균형이 가장 우수했다.{next_q_sentence} 따라서 본문과 부록의 최종 비교 표는 선택된 q={selected_q:.2f}를 기준으로 해석하는 것이 타당하다.",
        200: f"분석: cost split 기준으로 Selective Routing (q={selected_q:.2f})은 call rate {cost_selective['llm_call_rate_mean']:.4f}, cost ${cost_selective['cost_usd_mean']:.4f}, inference time {cost_selective['total_latency_ms_mean']/1000.0:.4f} s를 보였다. 이는 full-LLM upper bound의 cost ${cost_full_llm['cost_usd_mean']:.4f}와 time {cost_full_llm['total_latency_ms_mean']/1000.0:.4f} s보다 훨씬 낮으며, no-LLM baseline의 저비용 경로와 비교해선 추가 호출을 허용하는 대신 더 풍부한 경계 샘플 검토를 수행한다. 따라서 선택적 라우팅은 비용을 통제하면서도 운영상 필요한 보강 판단을 제공하는 실용적 절충안으로 해석할 수 있다.",
        201: "",
        320: "Figure 6는 운영 비용(Call Rate)과 탐지 안정성(Worst-Case Recall) 사이의 트레이드오프를 보여준다. 기존의 Full-LLM 방식은 높은 호출량과 비용을 요구하는 상한선이며, No-LLM 방식은 비용은 거의 없지만 경계 샘플 보호력이 제한적이다. 반면 본 연구의 선택적 라우팅은 q=0.40에서 평균 call rate 0.0825, cost $0.0088, time 42.7108 s 수준으로 cost split의 파레토 전선에 위치하여, 비용을 통제하면서도 실용적인 안정성을 확보한다.",
        391: "Table A2. Gray-zone Ratio and F1-score for individual base models across q-sweep.",
        311: f"본 연구의 실험 결과는 평균 AUC 극대화가 시스템의 진정한 가치가 아님을 보여준다. Table 2에서 UTAR는 AUC {utar_table2['AUC']}로 Avg. Ensemble의 {avg_ensemble_table2['AUC']}보다 높지 않지만, Ave. F1 {utar_table2['Ave. F1']}, PRR {utar_table2['PRR']}, Worst-Case Recall (P5) {utar_table2['Worst-Case Recall (P5)']}를 달성한다. 특히 ModernTCN의 Instability {moderntcn_table2['Instability']}와 비교하면 UTAR의 Instability는 {utar_table2['Instability']}로 더 낮아, shift 상황에서의 경계 관리와 운영 안정성 측면에서 우수한 trade-off를 보인다. 이는 UTAR의 핵심이 global score ranking을 키우는 것이 아니라, distribution shift 하에서 의사결정 임계치 주변 샘플을 더 안정적으로 처리하는 데 있음을 의미한다.",
        323: "핵심 메커니즘인 UTAR는 의사결정이 위험한 Gray-Zone을 명시적으로 식별하여 모델을 라우팅함으로써, 평균 지표에 가려진 recall 붕괴와 의사결정 불안정성을 완화하였다. 최종 선택된 q=0.40에서 UTAR는 Ave. F1 0.8835 ± 0.0073과 PRR 0.9807 ± 0.0110을 달성했고, cost split에서는 selective routing 비용이 $0.0088로 full-LLM upper bound의 $0.1073보다 약 91.8% 낮았다. 즉 본 프레임워크의 강점은 모든 지표의 절대 우세가 아니라, 시간적 분포 변화 하에서 성능-비용-안정성의 균형점을 실용적으로 확보한 데 있다.",
        342: "Table A1은 동일한 TDS 평가 프로토콜 아래에서 ModernTCN, AdapTable-inspired TTA, Cao-style invariant baseline, 그리고 UTAR를 비교한 결과를 요약한다. 본 부록 표의 목적은 제안 프레임워크가 단순 base ensemble 대비뿐 아니라 shift-aware baseline들과 비교해도 어떤 위치에 있는지를 정량적으로 보여주는 데 있다.",
        343: "재현성(Reproducibility)을 위해 Appendix B에는 현재 코드에서 사용하는 selective routing prompt의 실제 구조를 공개한다.",
        344: "B.1. System Prompt for Selective Routing in TEP",
        345: "[System Role] You are an expert Tennessee Eastman process engineer. Your task is to make the final routing-time decision for one ambiguous sample escalated near the UTAR decision boundary.",
        346: "[Output Rule] Return JSON only with {\"decision\": \"normal\"} or {\"decision\": \"anomaly\"}. No explanation, confidence, or extra keys are allowed.",
        347: "[Task Context] The sample was escalated because it lies near the UTAR boundary and is not a straightforward case. [Decision Objective] Missing a true anomaly is considered more costly than flagging a borderline anomaly.",
        348: "[Decision Policy] The prompt instructs the model to output anomaly when one strong anomaly signal or multiple moderate anomaly signals are jointly observed, especially when detector-side evidence and GraphAD+ structural evidence agree.",
        349: "[Guardrails] The model is explicitly told not to default to normal near the boundary and not to require perfect agreement across all detectors before choosing anomaly. In the actual deployed prompt, this system block is followed by three routing-style few-shot examples.",
        350: "B.2. Few-shot Examples and Serialized Context",
        351: "The actual prompt includes three routing-style few-shot examples: one normal case and two anomaly cases.",
        352: "Each example is serialized with the same header fields used at inference time, including utar_side, detector_anomaly_votes, max_detector, entropy_level, discrepancy_level, graphad_support, graphad_concentration, and sensor_coherence.",
        353: "The detection block then provides rf, xgb, ModernTCN, utar_base, ensemble_mean, temporal_weight, ensemble_entropy, and model_discrepancy.",
        354: "The GraphAD+ block provides graphad_score, top1_sensor, top1_score, top1_z, top1_trend, top1_fluct, top1_gap, topk_mean, candidate_sensors, candidate_scores, and candidate_topology.",
        355: "B.3. Prompt Semantics",
        356: "System Role는 모델에게 'TEP 경계 샘플의 최종 routing-time decision maker'라는 역할을 부여한다.",
        357: "Output Rule는 출력 형식을 anomaly/normal 이진 JSON으로 강제하여 후처리 파이프라인과 직접 연결되도록 한다.",
        358: "Task Context와 Decision Objective는 경계 샘플이라는 운영 맥락과 miss cost가 더 크다는 비용 구조를 함께 전달한다.",
        359: "Decision Policy는 detector vote, entropy, discrepancy, 그리고 GraphAD+ concentration/coherence를 조합해 anomaly 조건을 정의한다.",
        360: "Guardrails는 완벽한 합의가 없더라도 plausible anomaly evidence가 있으면 anomaly를 택하도록 유도하며, 경계 구간에서 default-normal 편향을 억제한다.",
        361: "Few-shot examples는 최종 fault label 예측이 아니라 routing decision(normal/anomaly) 자체를 학습시키는 예시로 구성된다.",
        362: "즉 Appendix B의 prompt는 일반적인 설명형 진단 프롬프트가 아니라, UTAR가 선별한 경계 샘플에 대해 최종 anomaly decision만 수행하도록 설계된 routing prompt이다.",
        363: "대표 프롬프트 예시는 appendix_b_prompt_example.txt와 table_b1_prompt_structure.csv, table_b2_prompt_variables.csv에 함께 저장된다.",
        364: "",
        365: "",
        366: "",
        367: "",
        368: "",
        369: "",
        370: "",
        371: "",
        372: "",
        373: "",
        383: "Table I. Diagnosis Results (Root-cause Ranking) on NDSS.",
        384: "Note: * indicates p < 0.05.",
        385: "NDSS 보조 분석에서 GraphAD+는 detection booster가 아니라 root-cause candidate ranking의 안정화 모듈로서 해석해야 하며, 성능 차이는 candidate ordering의 일관성 관점에서 읽는 것이 적절하다.",
        395: "내용: normal, transition, post_shift 구간 사이의 평균 이동, KS 통계량, KS reject 비율, 그리고 MMD를 통해 시프트 강도를 정량화한 통계 요약을 제시한다.",
        396: "Appendix E는 본문 4.2.1절의 temporal distribution shift 구성이 실제로 데이터 분포 차이를 유발했는지 보조적으로 확인하는 섹션이다. 여기서는 normal, transition, post_shift 구간 사이의 평균 이동, KS 통계량, KS reject 비율, 그리고 RBF-MMD를 함께 제시하여 설정한 shift 조건이 단순 라벨 분할이 아니라 실질적 분포 이동을 만들었음을 보여준다.",
        405: "(c) Auxiliary Hybrid Reranking: Root-Cause Identification",
        406: "최종 결과: 구조 점수와 LLM ranking score를 함께 고려하면, 통계적 중요도와 공정 맥락을 동시에 만족하는 root-cause candidate가 최상위로 정렬된다.",
        407: "이 시각화는 raw structural score만으로는 하위에 머물 수 있는 실제 원인 변수가, 구조 정보와 전문가형 추론의 결합을 통해 최종 상위 후보로 재정렬될 수 있음을 보여준다.",
    }
    for idx, text in replacements.items():
        _replace_paragraph_text(doc.paragraphs[idx], text)

    doc.save(TARGET_DOCX)
    _replace_images(TARGET_DOCX)
    doc = Document(TARGET_DOCX)
    _resize_inline_shape(doc, partname="/word/media/image1.png", width_inches=6.3, height_inches=5.6)
    doc.save(TARGET_DOCX)
    print(f"Updated document written to: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
