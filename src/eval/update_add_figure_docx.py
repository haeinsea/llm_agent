from __future__ import annotations

from copy import deepcopy
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
import argparse
from glob import glob
import json
import os
import shutil
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.append(str(PROJECT_ROOT))

from ndss.ndss_reasoning_cost import (
    auto_optimize_window,
    compute_ACT,
    graph_smooth_scores,
    graphad_scores_sliding,
    load_and_clean_csv,
    load_process_graph,
    reciprocal_rank_from_ranking,
)
from src.eval.update_reference_docx import (
    _build_appendix_a1_table,
    _build_appendix_g_seed_detail_table,
    _build_appendix_q_table,
    _build_table2,
    _build_table3,
    _build_table5,
    _build_table8,
)
from src.routing.selective_llm_eval import DEFAULT_Q, read_selected_q


APPENDIX_DIR = PROJECT_ROOT / "outputs" / "appendix"
DOCX_PATH = PROJECT_ROOT / "UTAR 논문(Reference)_260404)_add_figure.docx"
DOCX_BACKUP_PATH = PROJECT_ROOT / "UTAR 논문(Reference)_260404)_add_figure.backup_before_ndss_tables.docx"

_C1_PROC_GRAPH = None
_C1_ALPHA = None


def _set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)


def _clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _set_paragraph_text(paragraph, text: str, *, italic: bool = False) -> None:
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.italic = italic


def _remove_row_trailing_omission(row) -> None:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return
    for tag in ("w:gridAfter", "w:wAfter"):
        el = tr_pr.find(qn(tag))
        if el is not None:
            tr_pr.remove(el)


def _append_blank_cell(row) -> None:
    _remove_row_trailing_omission(row)
    new_tc = deepcopy(row._tr.tc_lst[-1])
    for node in new_tc.iter():
        if node.tag == qn("w:t"):
            node.text = ""
    row._tr.append(new_tc)


def _split_last_header_span(row) -> None:
    if not row._tr.tc_lst:
        return
    last_tc = row._tr.tc_lst[-1]
    tc_pr = last_tc.tcPr
    if tc_pr is None:
        return
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return
    span = int(grid_span.get(qn("w:val"), "1"))
    if span <= 1:
        return
    tc_w = tc_pr.find(qn("w:tcW"))
    split_width = None
    if tc_w is not None and tc_w.get(qn("w:w")):
        try:
            width = int(tc_w.get(qn("w:w")))
            split_width = max(width // span, 1)
            tc_w.set(qn("w:w"), str(split_width))
        except ValueError:
            split_width = None
    tc_pr.remove(grid_span)
    new_tc = deepcopy(last_tc)
    new_tc_pr = new_tc.find(qn("w:tcPr"))
    if new_tc_pr is not None:
        new_grid_span = new_tc_pr.find(qn("w:gridSpan"))
        if new_grid_span is not None:
            new_tc_pr.remove(new_grid_span)
        new_tc_w = new_tc_pr.find(qn("w:tcW"))
        if new_tc_w is not None and split_width is not None:
            new_tc_w.set(qn("w:w"), str(split_width))
    row._tr.append(new_tc)


def _write_dataframe_to_table(table, df: pd.DataFrame) -> None:
    text_df = df.fillna("").astype(str)
    while len(table.rows[0]._tr.tc_lst) < len(text_df.columns) and len(table.rows[0].cells) >= len(text_df.columns):
        _split_last_header_span(table.rows[0])
    while len(table.rows) < len(text_df) + 1:
        table.add_row()
    while len(table.rows) > len(text_df) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row in table.rows:
        while len(row.cells) < len(text_df.columns):
            _append_blank_cell(row)
    for j, col in enumerate(text_df.columns):
        _set_cell_text(table.rows[0].cells[j], col)
    for i, (_, row) in enumerate(text_df.iterrows(), start=1):
        for j, value in enumerate(row.tolist()):
            _set_cell_text(table.rows[i].cells[j], value)


def _insert_paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_table_after(paragraph, n_rows: int, n_cols: int, style_name: str | None = None) -> Table:
    container = paragraph._parent
    table = container.add_table(rows=n_rows, cols=n_cols, width=Inches(6.5))
    if style_name:
        table.style = style_name
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return Table(tbl, container)


def _replace_table_xml(target_table: Table, source_table: Table) -> None:
    new_tbl = deepcopy(source_table._tbl)
    target_table._tbl.addprevious(new_tbl)
    target_table._tbl.getparent().remove(target_table._tbl)


def _normalize_text(text: str) -> str:
    return " ".join((text or "").split())


def _find_paragraph(doc: Document, needle: str) -> tuple[int, Paragraph]:
    target = _normalize_text(needle)
    for idx, paragraph in enumerate(doc.paragraphs):
        if target in _normalize_text(paragraph.text):
            return idx, paragraph
    raise ValueError(f"Could not locate paragraph containing: {needle}")


def _format_metric(value: float) -> str:
    return f"{float(value):.3f}"


def _format_tokens(value: int | float) -> str:
    return f"{int(value):,}"


def _load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def _init_c1_worker(procmap_csv: str, alpha_graph: float) -> None:
    global _C1_PROC_GRAPH, _C1_ALPHA
    _C1_PROC_GRAPH = load_process_graph(procmap_csv)
    _C1_ALPHA = alpha_graph


def _eval_c1_scenario(args: tuple[str, list[str]]) -> dict[str, float | None]:
    csv_path, true_vars = args
    df = load_and_clean_csv(csv_path)
    best_w = auto_optimize_window(
        df,
        true_vars,
        g=None,
        w_z=1.0,
        w_trend=0.0,
        w_flat=0.0,
    )
    raw_scores = graphad_scores_sliding(
        df,
        win_len=best_w,
        g=None,
        alpha_graph=0.0,
        w_z=1.0,
        w_trend=0.0,
        w_flat=0.0,
    )
    if raw_scores is None or raw_scores.empty:
        return {
            "raw_top5": None,
            "raw_mrr": None,
            "smooth_top5": None,
            "smooth_mrr": None,
        }

    smooth_scores = graph_smooth_scores(raw_scores, _C1_PROC_GRAPH, alpha=_C1_ALPHA)
    raw_ranking = list(raw_scores.index)
    smooth_ranking = list(smooth_scores.index)
    return {
        "raw_top5": compute_ACT(raw_ranking[:5], true_vars),
        "raw_mrr": reciprocal_rank_from_ranking(raw_ranking, true_vars),
        "smooth_top5": compute_ACT(smooth_ranking[:5], true_vars),
        "smooth_mrr": reciprocal_rank_from_ranking(smooth_ranking, true_vars),
    }


def _parse_lambda_tuple(raw: str) -> tuple[float, float, float]:
    text = str(raw).strip().replace("[", "(").replace("]", ")")
    text = text.strip("()")
    parts = [float(part.strip()) for part in text.split(",") if part.strip()]
    if len(parts) != 3:
        raise ValueError(f"Invalid lambda tuple: {raw}")
    return (parts[0], parts[1], parts[2])


def _select_row_by_float(df: pd.DataFrame, column: str, value: float) -> pd.Series:
    matches = df[df[column].astype(float).round(6).eq(round(float(value), 6))]
    if matches.empty:
        raise ValueError(f"Could not find row where {column} == {value}")
    return matches.iloc[0]


def _select_row_by_lambda(df: pd.DataFrame, lambda_tuple: tuple[float, float, float]) -> pd.Series:
    if {"lambda_z", "lambda_tr", "lambda_fl"}.issubset(df.columns):
        matches = df[
            df["lambda_z"].astype(float).round(6).eq(round(lambda_tuple[0], 6))
            & df["lambda_tr"].astype(float).round(6).eq(round(lambda_tuple[1], 6))
            & df["lambda_fl"].astype(float).round(6).eq(round(lambda_tuple[2], 6))
        ]
        if not matches.empty:
            return matches.iloc[0]
    if "lambda_tuple" in df.columns:
        for _, row in df.iterrows():
            if _parse_lambda_tuple(row["lambda_tuple"]) == lambda_tuple:
                return row
    raise ValueError(f"Could not find lambda row for {lambda_tuple}")


def _build_table6_df(no_llm_summary_path: Path, cost_summary_csv: Path) -> pd.DataFrame:
    no_llm = _load_json(no_llm_summary_path)
    cost = pd.read_csv(cost_summary_csv)
    keep_models = ["gpt-4o-mini", "gpt-4.1", "gpt-5"]

    rows = [
        {
            "Mode": "No-LLM(Base)",
            "Model": "No-LLM(Base)",
            "Model_dup": "-",
            "K-Concord": _format_metric(no_llm["K-Concord"]),
            "SGR": _format_metric(no_llm["SGR"]),
            "Top1 Recall": _format_metric(no_llm["Top1_recall"]),
            "Top3 Recall": _format_metric(no_llm["Top3_recall"]),
            "Top5 Recall": _format_metric(no_llm["Top5_recall"]),
        }
    ]
    for model in keep_models:
        row = cost[cost["model"] == model].iloc[0]
        rows.append(
            {
                "Mode": "Hybrid (Graph AD + LLM Hybrid Reranking)",
                "Model": "Hybrid (Graph AD + LLM Hybrid Reranking)",
                "Model_dup": model,
                "K-Concord": _format_metric(row["K-Concord"]),
                "SGR": _format_metric(row["SGR"]),
                "Top1 Recall": _format_metric(row["Top1_recall"]),
                "Top3 Recall": _format_metric(row["Top3_recall"]),
                "Top5 Recall": _format_metric(row["Top5_recall"]),
            }
        )
    return pd.DataFrame(rows)


def _compute_candidate_ranking_baselines(alpha_summary_csv: Path, selected_alpha: float) -> pd.DataFrame:
    gt_df = pd.read_csv(PROJECT_ROOT / "data" / "ndss_attack_scenarios.csv")
    gt_map = {
        row.attack_id: [item.strip() for item in str(row.true_var).split(";") if item.strip()]
        for row in gt_df.itertuples()
    }
    files = sorted(glob(str(PROJECT_ROOT / "data" / "ndss_scenarios" / "*.csv")))
    tasks = []
    for csv_path in files:
        attack_id = Path(csv_path).stem
        true_vars = gt_map.get(attack_id, [])
        if true_vars:
            tasks.append((csv_path, true_vars))

    raw_top5: list[float] = []
    raw_mrr: list[float] = []
    smooth_top5: list[float] = []
    smooth_mrr: list[float] = []

    procmap_csv = str(PROJECT_ROOT / "data" / "NDSS_process_edges.csv")
    worker_count = min(8, os.cpu_count() or 1)
    try:
        executor = ProcessPoolExecutor(
            max_workers=worker_count,
            initializer=_init_c1_worker,
            initargs=(procmap_csv, selected_alpha),
        )
    except PermissionError:
        _init_c1_worker(procmap_csv, selected_alpha)
        executor = ThreadPoolExecutor(max_workers=worker_count)

    with executor as ex:
        for idx, result in enumerate(ex.map(_eval_c1_scenario, tasks), start=1):
            if result["raw_top5"] is not None:
                raw_top5.append(float(result["raw_top5"]))
                raw_mrr.append(float(result["raw_mrr"]))
                smooth_top5.append(float(result["smooth_top5"]))
                smooth_mrr.append(float(result["smooth_mrr"]))
            if idx % 50 == 0:
                print(f"Processed NDSS ranking baselines: {idx}/{len(tasks)} scenarios", flush=True)

    rows = []
    for name, top5_vals, mrr_vals in [
        ("Robust z-score", raw_top5, raw_mrr),
        ("Graph smoothing", smooth_top5, smooth_mrr),
    ]:
        top5 = np.asarray(top5_vals, dtype=float)
        mrr = np.asarray(mrr_vals, dtype=float)
        rows.append(
            {
                "Method": name,
                "Top-5 Hit": f"{top5.mean():.3f} ± {top5.std(ddof=1):.3f}",
                "MRR": f"{mrr.mean():.3f} ± {mrr.std(ddof=1):.3f}",
            }
        )
    alpha_summary = pd.read_csv(alpha_summary_csv)
    current_graphad = _select_row_by_float(alpha_summary, "alpha_graph", selected_alpha)
    rows.append(
        {
            "Method": "GraphAD+",
            "Top-5 Hit": f"{float(current_graphad['Top5_recall_mean']):.3f} ± {float(current_graphad['Top5_recall_std']):.3f}",
            "MRR": f"{float(current_graphad['MRR_mean']):.3f} ± {float(current_graphad['MRR_std']):.3f}",
        }
    )
    return pd.DataFrame(rows)


def _build_table_k1_df(no_llm_summary_path: Path, cost_summary_csv: Path) -> pd.DataFrame:
    no_llm = _load_json(no_llm_summary_path)
    cost = pd.read_csv(cost_summary_csv)
    order = [
        "gpt-4.1-nano",
        "gpt-4o-mini",
        "gpt-4.1-mini",
        "gpt-4.1",
        "gpt-4o",
        "gpt-5-nano",
        "gpt-5-mini",
        "gpt-5",
    ]

    rows = [
        {
            "Mode": "No-LLM",
            "Mode_dup": "-",
            "K-Concord": "-",
            "K-Concord_dup": _format_metric(no_llm["K-Concord"]),
            "SGR": _format_metric(no_llm["SGR"]),
            "Top1 Recall": _format_metric(no_llm["Top1_recall"]),
            "Top3 Recall": _format_metric(no_llm["Top3_recall"]),
            "Top5 Recall": _format_metric(no_llm["Top5_recall"]),
            "Prompt Tokens": "-",
            "Completion Tokens": "-",
            "Total Tokens": "-",
            "Cost (USD)": "-",
        }
    ]
    for model in order:
        row = cost[cost["model"] == model].iloc[0]
        rows.append(
            {
                "Mode": "Hybrid" if model == order[0] else "",
                "Mode_dup": model,
                "K-Concord": model,
                "K-Concord_dup": _format_metric(row["K-Concord"]),
                "SGR": _format_metric(row["SGR"]),
                "Top1 Recall": _format_metric(row["Top1_recall"]),
                "Top3 Recall": _format_metric(row["Top3_recall"]),
                "Top5 Recall": _format_metric(row["Top5_recall"]),
                "Prompt Tokens": _format_tokens(row["prompt_tokens"]),
                "Completion Tokens": _format_tokens(row["completion_tokens"]),
                "Total Tokens": _format_tokens(row["total_tokens"]),
                "Cost (USD)": f"{float(row['cost_usd']):.3f}",
            }
        )
    return pd.DataFrame(rows)


def _build_appendix_h_summary_df(
    lambda_detail_csv: Path,
    alpha_detail_csv: Path,
    beta_detail_csv: Path,
    selected_lambda: tuple[float, float, float],
    selected_alpha: float,
    selected_beta: float,
) -> pd.DataFrame:
    lambda_df = pd.read_csv(lambda_detail_csv)
    alpha_df = pd.read_csv(alpha_detail_csv)
    beta_df = pd.read_csv(beta_detail_csv)

    lambda_row = _select_row_by_lambda(lambda_df, selected_lambda)
    alpha_row = _select_row_by_float(alpha_df, "alpha_graph", selected_alpha)
    beta_row = _select_row_by_float(beta_df, "beta", selected_beta)

    rows = [
        {
            "Item": "Lambda weights (λz, λtr, λfl)",
            "Fixed conditions": "alpha fixed during lambda sweep",
            "Candidate grid": "36 simplex tuples (step 0.1, sum=1)",
            "Selected value": lambda_row["lambda_tuple"],
            "Key evidence": (
                f"Top1/3/5={float(lambda_row['Top1_recall_mean']):.3f}/"
                f"{float(lambda_row['Top3_recall_mean']):.3f}/"
                f"{float(lambda_row['Top5_recall_mean']):.3f}, "
                f"MRR={float(lambda_row['MRR_mean']):.3f}, "
                f"RankVar={float(lambda_row['ranking_variance']):.3f}."
            ),
        },
        {
            "Item": "Graph smoothing alpha",
            "Fixed conditions": f"lambda fixed at {selected_lambda}",
            "Candidate grid": "0.0 to 0.6 (step 0.1)",
            "Selected value": f"{float(alpha_row['alpha_graph']):.1f}",
            "Key evidence": (
                f"Top1/3/5={float(alpha_row['Top1_recall_mean']):.3f}/"
                f"{float(alpha_row['Top3_recall_mean']):.3f}/"
                f"{float(alpha_row['Top5_recall_mean']):.3f}, "
                f"MRR={float(alpha_row['MRR_mean']):.3f}, "
                f"RankVar={float(alpha_row['ranking_variance']):.3f}."
            ),
        },
        {
            "Item": "Auxiliary hybrid beta",
            "Fixed conditions": f"lambda={selected_lambda}, alpha={selected_alpha:.1f}, gpt-4o-mini",
            "Candidate grid": "0.0 / 0.2 / 0.4",
            "Selected value": f"{float(beta_row['beta']):.1f}",
            "Key evidence": (
                f"Top1/3/5={float(beta_row['Top1_recall']):.3f}/"
                f"{float(beta_row['Top3_recall']):.3f}/"
                f"{float(beta_row['Top5_recall']):.3f}, "
                f"MRR={float(beta_row['MRR']):.3f}, "
                f"RankVar={float(beta_row['ranking_variance']):.3f}."
            ),
        },
    ]
    return pd.DataFrame(rows)


def _write_csv_artifacts(table6_df: pd.DataFrame, table_c1_df: pd.DataFrame, table_k1_df: pd.DataFrame, appendix_h_df: pd.DataFrame) -> None:
    APPENDIX_DIR.mkdir(parents=True, exist_ok=True)
    table6_df.to_csv(APPENDIX_DIR / "table6_ndss_diagnostic_consistency.csv", index=False)
    table_c1_df.to_csv(APPENDIX_DIR / "table_c1_ndss_root_cause_ranking.csv", index=False)
    table_k1_df.to_csv(APPENDIX_DIR / "table_k1_ndss_root_cause_candidate_performance.csv", index=False)
    appendix_h_df.to_csv(APPENDIX_DIR / "table_h11_ndss_param_sensitivity_summary.csv", index=False)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip_compute", action="store_true")
    parser.add_argument("--lambda_z", type=float, default=0.1)
    parser.add_argument("--lambda_tr", type=float, default=0.7)
    parser.add_argument("--lambda_fl", type=float, default=0.2)
    parser.add_argument("--alpha_graph", type=float, required=True)
    parser.add_argument("--beta", type=float, required=True)
    parser.add_argument("--no_llm_summary", default=str(APPENDIX_DIR / "ndss_performance_summary_no_llm_selected.json"))
    parser.add_argument("--cost_summary_csv", default=str(APPENDIX_DIR / "ndss_cost_summary_selected.csv"))
    parser.add_argument("--lambda_detail_csv", default=str(APPENDIX_DIR / "table_h2_graphad_lambda_grid_detail_ndss.csv"))
    parser.add_argument("--alpha_detail_csv", default=str(APPENDIX_DIR / "table_h6_graphad_alpha_grid_detail_ndss.csv"))
    parser.add_argument("--alpha_summary_csv", default=str(APPENDIX_DIR / "table_h6_graphad_alpha_selection_ndss.csv"))
    parser.add_argument("--beta_detail_csv", default=str(APPENDIX_DIR / "table_h10_graphad_beta_grid_detail_ndss.csv"))
    args = parser.parse_args()

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    selected_lambda = (args.lambda_z, args.lambda_tr, args.lambda_fl)
    no_llm_summary_path = Path(args.no_llm_summary)
    cost_summary_csv = Path(args.cost_summary_csv)
    lambda_detail_csv = Path(args.lambda_detail_csv)
    alpha_detail_csv = Path(args.alpha_detail_csv)
    alpha_summary_csv = Path(args.alpha_summary_csv)
    beta_detail_csv = Path(args.beta_detail_csv)

    if not DOCX_BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, DOCX_BACKUP_PATH)

    if args.skip_compute:
        table6_df = pd.read_csv(APPENDIX_DIR / "table6_ndss_diagnostic_consistency.csv", dtype=str)
        table_c1_df = pd.read_csv(APPENDIX_DIR / "table_c1_ndss_root_cause_ranking.csv", dtype=str)
        table_k1_df = pd.read_csv(APPENDIX_DIR / "table_k1_ndss_root_cause_candidate_performance.csv", dtype=str)
        appendix_h_df = pd.read_csv(APPENDIX_DIR / "table_h11_ndss_param_sensitivity_summary.csv", dtype=str)
    else:
        table6_df = _build_table6_df(no_llm_summary_path=no_llm_summary_path, cost_summary_csv=cost_summary_csv)
        table_c1_df = _compute_candidate_ranking_baselines(
            alpha_summary_csv=alpha_summary_csv,
            selected_alpha=args.alpha_graph,
        )
        table_k1_df = _build_table_k1_df(no_llm_summary_path=no_llm_summary_path, cost_summary_csv=cost_summary_csv)
        appendix_h_df = _build_appendix_h_summary_df(
            lambda_detail_csv=lambda_detail_csv,
            alpha_detail_csv=alpha_detail_csv,
            beta_detail_csv=beta_detail_csv,
            selected_lambda=selected_lambda,
            selected_alpha=args.alpha_graph,
            selected_beta=args.beta,
        )
        _write_csv_artifacts(table6_df, table_c1_df, table_k1_df, appendix_h_df)

    table6_df = table6_df.copy()
    table6_df.columns = ["Mode", "Model", "Model", "K-Concord", "SGR", "Top1 Recall", "Top3 Recall", "Top5 Recall"]

    table_k1_df = table_k1_df.copy()
    table_k1_df.columns = ["Mode", "Mode", "K-Concord", "K-Concord", "SGR", "Top1 Recall", "Top3 Recall", "Top5 Recall", "Prompt Tokens", "Completion Tokens", "Total Tokens", "Cost (USD)"]

    doc = Document(DOCX_PATH)
    if DOCX_BACKUP_PATH.exists() and len(doc.tables) >= 19:
        backup_doc = Document(DOCX_BACKUP_PATH)
        _replace_table_xml(doc.tables[-2], backup_doc.tables[16])
        doc.save(DOCX_PATH)
        doc = Document(DOCX_PATH)

    selected_q = read_selected_q(DEFAULT_Q)
    _write_dataframe_to_table(doc.tables[1], _build_table2(selected_q))
    _write_dataframe_to_table(doc.tables[2], _build_table3())
    _write_dataframe_to_table(doc.tables[4], _build_table8())
    _write_dataframe_to_table(doc.tables[6], _build_table5(selected_q))
    _write_dataframe_to_table(doc.tables[10], _build_appendix_a1_table())
    _write_dataframe_to_table(doc.tables[12], _build_appendix_q_table())
    _write_dataframe_to_table(doc.tables[14], _build_appendix_g_seed_detail_table())

    _write_dataframe_to_table(doc.tables[5], table6_df)
    _write_dataframe_to_table(doc.tables[11], table_c1_df)
    _write_dataframe_to_table(doc.tables[-1], table_k1_df)

    _, appendix_h_anchor = _find_paragraph(
        doc,
        "즉 Process Graph Smoothing은 raw sensor data 전처리 제거/추가의 문제가 아니라",
    )
    summary_caption = "Summary of NDSS sensitivity evidence for lambda, alpha, and auxiliary beta."
    existing_caption_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if _normalize_text(paragraph.text) == _normalize_text(summary_caption):
            existing_caption_idx = idx
            break

    if existing_caption_idx is None:
        caption_paragraph = _insert_paragraph_after(appendix_h_anchor)
        _set_paragraph_text(caption_paragraph, summary_caption, italic=True)
        inserted_table = _insert_table_after(caption_paragraph, 1, len(appendix_h_df.columns), style_name=doc.tables[11].style.name)
        _write_dataframe_to_table(inserted_table, appendix_h_df)
    else:
        caption_paragraph = doc.paragraphs[existing_caption_idx]
        _set_paragraph_text(caption_paragraph, summary_caption, italic=True)
        anchor_tbl = None
        next_idx = existing_caption_idx + 1
        para_el = doc.paragraphs[existing_caption_idx]._p
        next_el = para_el.getnext()
        if next_el is not None and next_el.tag == qn("w:tbl"):
            anchor_tbl = Table(next_el, caption_paragraph._parent)
        if anchor_tbl is None:
            anchor_tbl = _insert_table_after(caption_paragraph, 1, len(appendix_h_df.columns), style_name=doc.tables[11].style.name)
        _write_dataframe_to_table(anchor_tbl, appendix_h_df)

    _set_paragraph_text(
        doc.paragraphs[124],
        "여기서 각 Score는 각각 순간 편차, 방향성 변화, 비정상 진동을 반영하며, 세 계수의 합은 1이 되도록 정규화된다. Appendix H에는 NDSS 기반 lambda, alpha, auxiliary beta 민감도 비교와 최종 선택 근거를 함께 요약하였다.",
    )
    _set_paragraph_text(
        doc.paragraphs[337],
        "Note: Top-5 Hit and MRR were computed on the same 286 NDSS scenarios; larger values indicate better root-cause ranking quality.",
    )
    _set_paragraph_text(
        doc.paragraphs[338],
        "Table C.1 compares three candidate-ranking variants on the identical 286 NDSS scenarios: a raw robust z-score baseline, a graph-smoothed ranking, and the current GraphAD+ setting. The purpose of this table is to show how structural smoothing and the final GraphAD+ fusion change Top-5 coverage and reciprocal-rank quality under the same evaluation protocol. The interpretation should focus on candidate ordering stability rather than standalone detection gain.",
    )

    doc.save(DOCX_PATH)
    print(f"Updated DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
